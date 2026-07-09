# app.py
import io
from datetime import datetime
import numpy as np
import pandas as pd
import streamlit as st
import matplotlib.pyplot as plt
import matplotlib
from matplotlib.font_manager import FontProperties

import colour
from colour import MSDS_CMFS

# Opcional: Savitzky-Golay si SciPy esta instalado
try:
    from scipy.signal import savgol_filter  # type: ignore
    _HAS_SCIPY = True
except Exception:
    _HAS_SCIPY = False

# Opcional: exportar reporte de rendimiento cuantico a Word
try:
    from docx import Document  # type: ignore
    from docx.shared import Pt, RGBColor  # type: ignore
    _HAS_DOCX = True
except Exception:
    _HAS_DOCX = False

# Opcional: exportar reporte de rendimiento cuantico a PDF
try:
    from reportlab.lib.pagesizes import letter  # type: ignore
    from reportlab.lib.styles import getSampleStyleSheet  # type: ignore
    from reportlab.lib import colors as rl_colors  # type: ignore
    from reportlab.lib.units import inch as rl_inch  # type: ignore
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle  # type: ignore
    _HAS_REPORTLAB = True
except Exception:
    _HAS_REPORTLAB = False

# ----------------- Config -----------------
st.set_page_config(layout="wide", page_title="SpectraLab Toolkit")
matplotlib.rcParams.update({'font.size': 12})
DIALOG_DECORATOR = getattr(st, 'dialog', None) or getattr(st, 'experimental_dialog', None)
HAS_DIALOG = callable(DIALOG_DECORATOR)

if "active_page" not in st.session_state:
    st.session_state["active_page"] = "Inicio"

# Hover suave (tinte rojo claro + sombra) para todos los botones de la app
st.markdown(
    """
    <style>
    div[data-testid="stButton"] > button {
        transition: background-color .15s ease, border-color .15s ease,
                    color .15s ease, box-shadow .15s ease;
    }
    div[data-testid="stButton"] > button:hover {
        background-color: #fff1f1 !important;
        border-color: #f87171 !important;
        color: #b91c1c !important;
        box-shadow: 0 2px 6px rgba(248, 113, 113, 0.25);
    }
    </style>
    """,
    unsafe_allow_html=True,
)



def go_to_page(page_name: str):
    st.session_state["active_page"] = page_name
    st.rerun()


# ============================================================
# Utilidades compartidas
# (unica implementacion, usada por todas las paginas: antes existian
# 2-3 copias casi identicas de cada una de estas funciones)
# ============================================================

def uploaded_to_bytes(uploaded):
    if uploaded is None:
        return b""
    try:
        return uploaded.getvalue()
    except Exception:
        try:
            uploaded.seek(0)
        except Exception:
            pass
        return uploaded.read()


def read_csv_flexible_bytes(raw: bytes) -> pd.DataFrame:
    if not raw:
        raise ValueError("El archivo esta vacio o no se pudo leer.")
    trials = [
        (";", ","), (";", "."),
        (",", ","), (",", "."),
        (None, ","), (None, "."),
    ]
    last_err = None
    for sep, dec in trials:
        try:
            bio = io.BytesIO(raw)
            df = pd.read_csv(bio, sep=sep, engine="python" if sep is None else "c", decimal=dec)
            if df.shape[1] >= 2:
                return df
        except Exception as e:
            last_err = e
    try:
        bio = io.BytesIO(raw)
        df = pd.read_csv(bio, sep=r"\s+", engine="python")
        if df.shape[1] >= 2:
            return df
    except Exception as e:
        last_err = e
    raise ValueError(f"No se pudo leer el CSV (separador/decimal). Ultimo error: {last_err}")


@st.cache_data(show_spinner=False)
def load_csv_df(raw: bytes) -> pd.DataFrame:
    return read_csv_flexible_bytes(raw)


@st.cache_data(show_spinner=False)
def list_excel_sheets(raw: bytes):
    return pd.ExcelFile(io.BytesIO(raw)).sheet_names


@st.cache_data(show_spinner=False)
def load_excel_sheet(raw: bytes, sheet: str) -> pd.DataFrame:
    return pd.read_excel(io.BytesIO(raw), sheet_name=sheet)


def to_numeric_series(s: pd.Series):
    return pd.to_numeric(s.astype(str).str.replace(",", "."), errors="coerce")


_WL_HINTS = ["wavelength", "wavelength (nm)", "lambda", "wl", "nm", "wave"]
_INT_HINTS = [
    "intensity", "intensity (a.u.)", "a.u.", "au", "counts", "cps", "signal",
    "emission", "fluorescence", "fluorescencia", "abs", "excitation",
]


def _norm_colname(c):
    return str(c).strip().lower().replace("\ufeff", "").replace("\u00b5", "u")


def guess_columns(df: pd.DataFrame):
    """Adivina columna de longitud de onda e intensidad por nombre, con
    fallback a la fraccion de valores numericos si el nombre no ayuda."""
    cols = list(df.columns)
    if not cols:
        return None, None
    if len(cols) < 2:
        return cols[0], cols[0]

    norm = [_norm_colname(c) for c in cols]
    wl_idx, int_idx = None, None
    for i, n in enumerate(norm):
        if any(h in n for h in _WL_HINTS):
            wl_idx = i
            break
    for i, n in enumerate(norm):
        if any(h in n for h in _INT_HINTS):
            int_idx = i
            break

    if wl_idx is None or int_idx is None or wl_idx == int_idx:
        numeric_scores = []
        for i, c in enumerate(cols):
            s = to_numeric_series(df[c])
            numeric_scores.append((i, float(s.notna().mean())))
        numeric_scores.sort(key=lambda x: x[1], reverse=True)
        if wl_idx is None and numeric_scores:
            wl_idx = numeric_scores[0][0]
        if int_idx is None and len(numeric_scores) > 1:
            int_idx = numeric_scores[1][0]
        if wl_idx == int_idx and len(numeric_scores) > 1:
            int_idx = numeric_scores[1][0]

    return (cols[wl_idx] if wl_idx is not None else cols[0],
            cols[int_idx] if int_idx is not None else cols[-1])


def spectrum_metrics(wl, intensity):
    wl = np.asarray(wl, dtype=float)
    intensity = np.asarray(intensity, dtype=float)
    integrate = getattr(np, "trapezoid", None) or getattr(np, "trapz")
    peak_idx = int(np.nanargmax(intensity))
    peak_wl = float(wl[peak_idx])
    peak_intensity = float(intensity[peak_idx])
    area = float(integrate(intensity, wl))
    half = peak_intensity / 2.0
    above = np.where(intensity >= half)[0]
    fwhm = np.nan
    if len(above) >= 2:
        fwhm = float(wl[int(above[-1])] - wl[int(above[0])])
    return peak_wl, peak_intensity, area, fwhm


def filter_and_normalize(wl_raw, int_raw, wl_min, wl_max, normalize="None"):
    """Filtra por rango y normaliza sin interpolar a una malla (usado en el
    visor, donde se quiere respetar el muestreo original del instrumento)."""
    data = pd.DataFrame({
        "wavelength": to_numeric_series(wl_raw),
        "intensity": to_numeric_series(int_raw),
    }).dropna()
    data = data[(data["wavelength"] >= wl_min) & (data["wavelength"] <= wl_max)].copy()
    if data.empty:
        raise ValueError("No hay datos dentro del rango seleccionado.")
    data = data.sort_values("wavelength").groupby("wavelength", as_index=False)["intensity"].mean()
    wl = data["wavelength"].to_numpy(dtype=float)
    intensity = data["intensity"].to_numpy(dtype=float)

    integrate = getattr(np, "trapezoid", None) or getattr(np, "trapz")
    if normalize == "Max = 1":
        max_val = float(np.nanmax(np.abs(intensity)))
        if max_val > 0:
            intensity = intensity / max_val
    elif normalize == "Area = 1":
        area_abs = float(integrate(np.abs(intensity), wl))
        if area_abs > 0:
            intensity = intensity / area_abs
    return wl, intensity


@st.cache_data(show_spinner=False)
def cmfs_arrays():
    cmfs = MSDS_CMFS['CIE 1931 2 Degree Standard Observer']
    if hasattr(cmfs, "wavelengths"):
        wls = np.asarray(cmfs.wavelengths, dtype=float)
    elif hasattr(cmfs, "domain"):
        wls = np.asarray(cmfs.domain, dtype=float)
    else:
        wls = np.asarray(cmfs.index, dtype=float)
    vals = np.asarray(cmfs.values, dtype=float)
    if vals.shape[0] == 3 and vals.shape[1] == wls.shape[0]:
        vals = vals.T
    return wls, vals[:, 0], vals[:, 1], vals[:, 2]


CMF_WLS, CMF_X, CMF_Y, CMF_Z = cmfs_arrays()
CMF_MIN, CMF_MAX = float(np.min(CMF_WLS)), float(np.max(CMF_WLS))


@st.cache_data(show_spinner=False)
def locus_arrays(wl_start=380, wl_end=780):
    m = (CMF_WLS >= wl_start) & (CMF_WLS <= wl_end)
    wls = CMF_WLS[m]
    xbar, ybar, zbar = CMF_X[m], CMF_Y[m], CMF_Z[m]
    den = xbar + ybar + zbar
    den = np.where(den == 0, np.nan, den)
    x = xbar / den
    y = ybar / den
    xy = np.column_stack([x, y])
    ok = np.isfinite(xy).all(axis=1)
    return wls[ok].astype(float), xy[ok]


LOCUS_WLS_F, LOCUS_XY = locus_arrays(380, 780)


def _cross2(a, b):
    return a[0] * b[1] - a[1] * b[0]


def _ray_segment_intersection(w, v, p0, p1, eps=1e-12):
    s = p1 - p0
    rxs = _cross2(v, s)
    if abs(rxs) < eps:
        return None
    q_p = p0 - w
    t = _cross2(q_p, s) / rxs
    u = _cross2(q_p, v) / rxs
    if t >= 0 and 0 <= u <= 1:
        return t, u
    return None


def dominant_wavelength_and_purity(x, y, white_point=(0.3333, 0.3333)):
    w = np.array([white_point[0], white_point[1]], dtype=float)
    p = np.array([x, y], dtype=float)
    v = p - w
    if not np.isfinite(v).all() or (abs(v[0]) < 1e-15 and abs(v[1]) < 1e-15):
        return np.nan, np.nan, "Indefinido"

    best = None  # (t, kind, i, u, q)
    for i in range(len(LOCUS_XY) - 1):
        p0, p1 = LOCUS_XY[i], LOCUS_XY[i + 1]
        hit = _ray_segment_intersection(w, v, p0, p1)
        if hit is None:
            continue
        t, u = hit
        if best is None or t < best[0]:
            best = (t, "locus", i, u, w + t * v)

    p0, p1 = LOCUS_XY[-1], LOCUS_XY[0]
    hit = _ray_segment_intersection(w, v, p0, p1)
    if hit is not None:
        t, u = hit
        if best is None or t < best[0]:
            best = (t, "purple", -1, u, w + t * v)

    if best is None:
        return np.nan, np.nan, "Sin interseccion"

    _, kind, i, u, q = best
    d_wp_p = float(np.linalg.norm(p - w))
    d_wp_q = float(np.linalg.norm(q - w))
    purity = (d_wp_p / d_wp_q) * 100.0 if d_wp_q > 0 else np.nan

    if kind == "locus":
        wl0, wl1 = float(LOCUS_WLS_F[i]), float(LOCUS_WLS_F[i + 1])
        wl = wl0 + u * (wl1 - wl0)
        return wl, purity, "Dominante"

    # longitud de onda complementaria (rayo opuesto)
    v2 = -v
    best2 = None
    for j in range(len(LOCUS_XY) - 1):
        p0, p1 = LOCUS_XY[j], LOCUS_XY[j + 1]
        hit2 = _ray_segment_intersection(w, v2, p0, p1)
        if hit2 is None:
            continue
        t2, u2 = hit2
        if best2 is None or t2 < best2[0]:
            best2 = (t2, j, u2)

    if best2 is None:
        return np.nan, purity, "Purpura (sin complementaria)"

    _, j, u2 = best2
    wl0, wl1 = float(LOCUS_WLS_F[j]), float(LOCUS_WLS_F[j + 1])
    wl_comp = wl0 + u2 * (wl1 - wl0)
    return wl_comp, purity, "Complementaria"


def spectrum_to_xy(wl_grid, intensity_grid):
    xbar = np.interp(wl_grid, CMF_WLS, CMF_X)
    ybar = np.interp(wl_grid, CMF_WLS, CMF_Y)
    zbar = np.interp(wl_grid, CMF_WLS, CMF_Z)

    integrate = getattr(np, "trapezoid", None) or getattr(np, "trapz")
    X = float(integrate(intensity_grid * xbar, wl_grid))
    Y = float(integrate(intensity_grid * ybar, wl_grid))
    Z = float(integrate(intensity_grid * zbar, wl_grid))

    S = X + Y + Z
    if not np.isfinite(S) or S <= 0:
        raise ValueError("Valores XYZ invalidos (las intensidades son casi cero en el rango seleccionado).")
    return X / S, Y / S


def preprocess_spectrum(df, wl_col, int_col,
                        wl_min_local, wl_max_local, interp_interval_local,
                        clip_negative=False,
                        baseline_subtract_min=False,
                        smooth_method="None",
                        smooth_window=11,
                        smooth_poly=3,
                        normalize_mode="None"):
    out = pd.DataFrame()
    out["wavelength"] = to_numeric_series(df[wl_col])
    out["intensity"] = to_numeric_series(df[int_col])
    out = out.dropna()
    out = out[(out["wavelength"] >= wl_min_local) & (out["wavelength"] <= wl_max_local)].copy()
    if out.empty:
        raise ValueError("No hay datos dentro del rango de longitud de onda seleccionado.")
    out = out.sort_values("wavelength")
    out = out.groupby("wavelength", as_index=False)["intensity"].mean()

    if baseline_subtract_min:
        out["intensity"] = out["intensity"] - float(out["intensity"].min())

    if clip_negative:
        out["intensity"] = out["intensity"].clip(lower=0)

    wl_grid = np.arange(wl_min_local, wl_max_local + 1e-9, interp_interval_local, dtype=float)
    intensity_grid = np.interp(
        wl_grid,
        out["wavelength"].values.astype(float),
        out["intensity"].values.astype(float),
        left=0.0, right=0.0
    )

    if smooth_method == "Moving average":
        w = int(max(3, smooth_window))
        if w % 2 == 0:
            w += 1
        kernel = np.ones(w, dtype=float) / w
        intensity_grid = np.convolve(intensity_grid, kernel, mode="same")
    elif smooth_method == "Savitzky-Golay" and _HAS_SCIPY:
        n = int(len(intensity_grid))
        w = min(int(max(5, smooth_window)), n if n % 2 == 1 else n - 1)
        if w % 2 == 0:
            w -= 1
        p = int(max(2, smooth_poly))
        if w >= 3 and p >= w:
            p = w - 1
        if w >= 3 and p < w:
            intensity_grid = savgol_filter(intensity_grid, window_length=w, polyorder=p, mode="interp")

    integrate = getattr(np, "trapezoid", None) or getattr(np, "trapz")
    if normalize_mode == "Max = 1":
        m = float(np.max(np.abs(intensity_grid)))
        if m > 0:
            intensity_grid = intensity_grid / m
    elif normalize_mode == "Area = 1":
        area = float(integrate(np.abs(intensity_grid), wl_grid))
        if area > 0:
            intensity_grid = intensity_grid / area

    if float(np.max(np.abs(intensity_grid))) <= 0:
        raise ValueError("Intensidad cero en el rango seleccionado (revisa columnas y rango).")

    return wl_grid, intensity_grid


def _safe_mpl_text(s: str) -> str:
    """Evita fallos de Matplotlib causados por mathtext invalido.

    Matplotlib intenta interpretar mathtext cuando ve '$'. Si el usuario
    escribe un solo '$' suelto, tight_layout puede fallar.
    """
    s = "" if s is None else str(s)
    if s.count("$") % 2 == 1:
        s = s.replace("$", "")
    return s


def show_and_close(fig):
    """st.pyplot + cierre explicito de la figura para evitar fugas de
    memoria en sesiones largas con muchos reruns."""
    st.pyplot(fig)
    plt.close(fig)


# ------------------------------------------------------------------
# Dataset de referencia embebido: espectro de emision real de sulfato de
# quinina en H2SO4 0.5 M, excitacion 310 nm, Phi_ref = 0.546 (Eaton, 1988).
# Datos digitalizados de PhotochemCAD 2.1a (Du et al. 1998; Dixon et al. 2005),
# reprocesados por Scott Prahl (Oregon Medical Laser Center). Se incluyen tal
# cual para que el usuario pueda usarlos como referencia real sin tener que
# descargarlos y volverlos a subir.
# Fuente: https://omlc.org/spectra/PhotochemCAD/html/081.html
# ------------------------------------------------------------------
_QUININE_SULFATE_EMS_TXT = """320	2281
320.5	2181
321	2175
321.5	2133
322	2121
322.5	2030
323	2084
323.5	2080
324	1973
324.5	2075
325	2125
325.5	2131
326	2124
326.5	2075
327	2091
327.5	2031
328	1953
328.5	2044
329	2184
329.5	1955
330	2073
330.5	1997
331	1981
331.5	1991
332	2019
332.5	2037
333	2071
333.5	1949
334	2036
334.5	2016
335	1977
335.5	2027
336	2020
336.5	2068
337	2039
337.5	2201
338	2251
338.5	2384
339	2440
339.5	2927
340	3059
340.5	3699
341	4179
341.5	5055
342	5745
342.5	6684
343	7479
343.5	8360
344	9407
344.5	9840
345	10201
345.5	10390
346	10544
346.5	10122
347	9839
347.5	9079
348	8229
348.5	7116
349	6263
349.5	5233
350	4554
350.5	3830
351	3333
351.5	2705
352	2610
352.5	2399
353	2221
353.5	2253
354	2097
354.5	2263
355	2218
355.5	2240
356	2158
356.5	2242
357	2144
357.5	2198
358	2337
358.5	2235
359	2227
359.5	2333
360	2356
360.5	2563
361	2444
361.5	2683
362	2611
362.5	2863
363	2955
363.5	3008
364	3142
364.5	3412
365	3549
365.5	3837
366	4128
366.5	4329
367	4667
367.5	5038
368	5238
368.5	5752
369	5946
369.5	6745
370	7139
370.5	7651
371	8374
371.5	9226
372	10062
372.5	11082
373	11862
373.5	12914
374	13706
374.5	15330
375	16082
375.5	18339
376	19429
376.5	21674
377	23419
377.5	25898
378	27814
378.5	30838
379	32688
379.5	36923
380	38961
380.5	42568
381	45621
381.5	50937
382	53543
382.5	59270
383	63413
383.5	70228
384	74002
384.5	80896
385	85941
385.5	95808
386	100075
386.5	109297
387	115835
387.5	127700
388	132832
388.5	144096
389	152316
389.5	165122
390	172334
390.5	184146
391	191923
391.5	206468
392	215380
392.5	229522
393	238909
393.5	256452
394	264660
394.5	280807
395	292369
395.5	310903
396	321926
396.5	337232
397	350524
397.5	370166
398	381733
398.5	400767
399	413848
399.5	437790
400	451603
400.5	471147
401	483912
401.5	511461
402	523055
402.5	545215
403	560394
403.5	589495
404	604237
404.5	627045
405	643460
405.5	678311
406	692439
406.5	722875
407	742917
407.5	778883
408	792766
408.5	821049
409	846161
409.5	880305
410	895378
410.5	926545
411	947605
411.5	980451
412	995351
412.5	1021701
413	1041357
413.5	1074101
414	1088811
414.5	1113170
415	1133364
415.5	1163363
416	1178283
416.5	1205510
417	1224503
417.5	1258939
418	1268136
418.5	1294770
419	1319152
419.5	1345360
420	1358315
420.5	1382288
421	1407324
421.5	1431855
422	1446909
422.5	1473928
423	1489268
423.5	1516672
424	1532743
424.5	1555275
425	1575196
425.5	1593824
426	1617423
426.5	1640678
427	1654585
427.5	1679729
428	1691094
428.5	1716570
429	1728504
429.5	1757626
430	1767316
430.5	1786297
431	1803027
431.5	1824190
432	1833171
432.5	1849754
433	1861380
433.5	1885376
434	1898917
434.5	1908920
435	1919955
435.5	1935792
436	1940502
436.5	1956126
437	1962567
437.5	1973844
438	1981945
438.5	1991856
439	2001829
439.5	2006503
440	2022301
440.5	2026296
441	2028185
441.5	2043288
442	2047012
442.5	2055358
443	2063902
443.5	2071054
444	2075214
444.5	2077171
445	2087976
445.5	2089010
446	2094813
446.5	2095598
447	2100185
447.5	2106433
448	2106650
448.5	2104651
449	2112311
449.5	2110937
450	2111394
450.5	2109571
451	2114415
451.5	2111157
452	2106323
452.5	2113838
453	2106578
453.5	2101827
454	2106674
454.5	2099943
455	2099573
455.5	2091466
456	2087621
456.5	2088939
457	2082004
457.5	2076582
458	2070829
458.5	2062576
459	2056803
459.5	2042820
460	2047611
460.5	2038843
461	2035336
461.5	2025322
462	2023044
462.5	2017491
463	2004979
463.5	1996031
464	1990110
464.5	1983785
465	1979172
465.5	1964159
466	1967267
466.5	1953132
467	1947229
467.5	1936273
468	1931617
468.5	1918285
469	1910536
469.5	1900605
470	1895311
470.5	1886174
471	1878382
471.5	1863258
472	1861710
472.5	1848459
473	1839928
473.5	1828934
474	1822824
474.5	1806764
475	1798805
475.5	1782589
476	1780090
476.5	1763788
477	1753484
477.5	1730645
478	1729738
478.5	1706604
479	1692376
479.5	1678256
480	1670853
480.5	1653553
481	1647147
481.5	1627837
482	1618623
482.5	1606353
483	1592124
483.5	1572231
484	1568348
484.5	1553255
485	1544612
485.5	1529931
486	1519791
486.5	1506386
487	1495218
487.5	1482477
488	1473545
488.5	1459846
489	1444732
489.5	1431975
490	1426928
490.5	1409764
491	1402359
491.5	1384869
492	1379777
492.5	1361501
493	1351411
493.5	1341712
494	1331609
494.5	1316615
495	1309431
495.5	1289553
496	1285405
496.5	1269725
497	1257902
497.5	1244259
498	1235220
498.5	1223225
499	1212461
499.5	1197916
500	1189137
500.5	1174565
501	1164506
501.5	1148754
502	1141238
502.5	1124609
503	1118343
503.5	1097335
504	1091489
504.5	1075798
505	1066900
505.5	1051580
506	1043887
506.5	1030209
507	1019606
507.5	1004322
508	997564
508.5	986826
509	976207
509.5	960196
510	960231
510.5	948312
511	935121
511.5	922812
512	915821
512.5	905655
513	896099
513.5	882514
514	877670
514.5	866843
515	859401
515.5	845016
516	841539
516.5	831188
517	821887
517.5	808601
518	806285
518.5	794342
519	786987
519.5	776193
520	768871
520.5	756820
521	751145
521.5	739840
522	736959
522.5	724092
523	716491
523.5	706362
524	703370
524.5	690815
525	682289
525.5	673302
526	664728
526.5	658304
527	651014
527.5	636870
528	637305
528.5	622699
529	616230
529.5	605065
530	602805
530.5	593068
531	585570
531.5	575988
532	570078
532.5	564199
533	556587
533.5	547866
534	542272
534.5	536734
535	530645
535.5	521050
536	516767
536.5	510500
537	505027
537.5	495899
538	494286
538.5	486969
539	481882
539.5	474191
540	468990
540.5	462136
541	458472
541.5	450963
542	447738
542.5	441028
543	436851
543.5	426921
544	424248
544.5	418841
545	415143
545.5	407036
546	402986
546.5	396425
547	393482
547.5	387593
548	382985
548.5	376405
549	375580
549.5	367464
550	363949
550.5	359827
551	354293
551.5	348503
552	344013
552.5	339187
553	333571
553.5	329949
554	326565
554.5	320257
555	316266
555.5	312021
556	309111
556.5	304336
557	302464
557.5	297034
558	293749
558.5	289562
559	285592
559.5	281347
560	279207
560.5	275338
561	272477
561.5	267178
562	264620
562.5	261740
563	258829
563.5	253167
564	251302
564.5	247426
565	245912
565.5	242308
566	237949
566.5	235529
567	232047
567.5	227743
568	227393
568.5	222848
569	222440
569.5	217347
570	215524
570.5	212062
571	210004
571.5	206347
572	203422
572.5	200171
573	198311
573.5	195473
574	195370
574.5	191522
575	188480
575.5	186254
576	183891
576.5	181109
577	178900
577.5	176202
578	174606
578.5	169736
579	169341
579.5	165173
580	164304
580.5	162826
581	161116
581.5	157138
582	155870
582.5	154549
583	152869
583.5	149688
584	149262
584.5	146467
585	144038
585.5	142865
586	140846
586.5	138827
587	136914
587.5	134659
588	133524
588.5	130691
589	130342
589.5	128304
590	125502
590.5	125953
591	124360
591.5	121478
592	120564
592.5	119129
593	116378
593.5	116478
594	114275
594.5	113511
595	111401
595.5	110368
596	108451
596.5	106886
597	105675
597.5	103110
598	103232
598.5	100672
599	100506
599.5	99162
600	96342
"""

_qs_df = pd.read_csv(io.StringIO(_QUININE_SULFATE_EMS_TXT), sep="\t", header=None, names=["wavelength", "intensity"])
QUININE_SULFATE_EXAMPLE = {
    "option_label": "Usar dato real (sulfato de quinina, PhotochemCAD)",
    "wl": _qs_df["wavelength"].to_numpy(dtype=float),
    "intensity": _qs_df["intensity"].to_numpy(dtype=float),
    "citation": (
        "Sulfato de quinina en H\u2082SO\u2084 0.5 M, excitacion 310 nm, \u03a6ref = 0.546 (Eaton, 1988). "
        "Espectro digitalizado: PhotochemCAD 2.1a (Du et al. 1998; Dixon et al. 2005), "
        "reprocesado por S. Prahl, Oregon Medical Laser Center."
    ),
    "url": "https://omlc.org/spectra/PhotochemCAD/html/081.html",
}


def _render_spectrum_result(wl, intensity, color="#1f77b4"):
    """Grafica un espectro y muestra/retorna su area integrada. Compartido
    entre la carga de archivo propio y el dataset de referencia embebido."""
    peak_wl, peak_intensity, area, fwhm = spectrum_metrics(wl, intensity)

    fig, ax = plt.subplots(figsize=(4, 2.0))
    ax.plot(wl, intensity, color=color, linewidth=1.6)
    ax.fill_between(wl, intensity, color=color, alpha=0.12)
    ax.axvline(peak_wl, color=color, linestyle="--", linewidth=0.8, alpha=0.6)
    ax.set_xlabel("Longitud de onda (nm)", fontsize=8)
    ax.set_ylabel("Intensidad", fontsize=8)
    ax.tick_params(labelsize=7)
    ax.grid(alpha=0.2)
    show_and_close(fig)

    m1, m2, m3 = st.columns(3)
    m1.metric("Area", f"{area:.4g}")
    m2.metric("Pico", f"{peak_wl:.1f} nm")
    m3.metric("FWHM", f"{fwhm:.1f} nm")
    return area



# ------------------------------------------------------------------
# Dataset de referencia embebido: coeficiente de extincion molar real de
# sulfato de quinina en H2SO4 0.5 M (Irvin & Irvin, 1948; digitalizado por
# PhotochemCAD 2.1a / Scott Prahl). Rango 255.5-420 nm: cubre toda la banda
# de absorcion real; se omite el resto (420-720 nm) porque son valores de
# ruido de fondo cercanos a cero, sin senal de absorcion real.
# Fuente: https://omlc.org/spectra/PhotochemCAD/html/081.html
# ------------------------------------------------------------------
_QUININE_SULFATE_ABS_TXT = """255.5	13372
255.75	13080
256	12992
256.25	13107
256.5	12924
256.75	12870
257	12852
257.25	12786
257.5	12621
257.75	12359
258	12135
258.25	11970
258.5	11734
258.75	11357
259	10790
259.25	10368
259.5	9834
259.75	9213
260	8636
260.25	8051
260.5	7466
260.75	6944
261	6421
261.25	5949
261.5	5448
261.75	5036
262	4607
262.25	4215
262.5	3867
262.75	3520
263	3234
263.25	2963
263.5	2717
263.75	2481
264	2283
264.25	2090
264.5	1917
264.75	1772
265	1626
265.25	1512
265.5	1401
265.75	1306
266	1226
266.25	1148
266.5	1075
266.75	1020
267	962
267.25	934
267.5	894
267.75	858
268	827
268.25	802
268.5	782
268.75	768
269	753
269.25	743
269.5	737
269.75	728
270	723
270.25	725
270.5	721
270.75	720
271	719
271.25	720
271.5	726
271.75	717
272	734
272.25	738
272.5	750
272.75	759
273	765
273.25	769
273.5	776
273.75	789
274	795
274.25	803
274.5	809
274.75	808
275	838
275.25	851
275.5	861
275.75	879
276	888
276.25	906
276.5	912
276.75	928
277	939
277.25	950
277.5	958
277.75	989
278	992
278.25	998
278.5	1007
278.75	1043
279	1049
279.25	1059
279.5	1073
279.75	1085
280	1096
280.25	1117
280.5	1132
280.75	1145
281	1165
281.25	1184
281.5	1201
281.75	1228
282	1232
282.25	1250
282.5	1269
282.75	1274
283	1298
283.25	1315
283.5	1341
283.75	1355
284	1372
284.25	1397
284.5	1412
284.75	1429
285	1440
285.25	1466
285.5	1487
285.75	1502
286	1520
286.25	1546
286.5	1563
286.75	1581
287	1608
287.25	1635
287.5	1652
287.75	1674
288	1685
288.25	1721
288.5	1737
288.75	1760
289	1784
289.25	1803
289.5	1835
289.75	1859
290	1877
290.25	1903
290.5	1931
290.75	1947
291	1973
291.25	2008
291.5	2028
291.75	2061
292	2096
292.25	2113
292.5	2141
292.75	2161
293	2201
293.25	2221
293.5	2258
293.75	2270
294	2301
294.25	2324
294.5	2354
294.75	2365
295	2394
295.25	2419
295.5	2460
295.75	2491
296	2532
296.25	2542
296.5	2579
296.75	2602
297	2643
297.25	2687
297.5	2719
297.75	2742
298	2773
298.25	2820
298.5	2863
298.75	2875
299	2913
299.25	2955
299.5	2978
299.75	3008
300	3042
300.25	3063
300.5	3095
300.75	3157
301	3176
301.25	3205
301.5	3255
301.75	3286
302	3321
302.25	3351
302.5	3381
302.75	3410
303	3449
303.25	3469
303.5	3498
303.75	3539
304	3560
304.25	3600
304.5	3626
304.75	3676
305	3698
305.25	3722
305.5	3738
305.75	3783
306	3820
306.25	3833
306.5	3877
306.75	3904
307	3941
307.25	3978
307.5	3992
307.75	4021
308	4035
308.25	4081
308.5	4117
308.75	4126
309	4159
309.25	4181
309.5	4199
309.75	4113
310	4155
310.25	4157
310.5	4196
310.75	4259
311	4259
311.25	4264
311.5	4301
311.75	4343
312	4331
312.25	4371
312.5	4377
312.75	4371
313	4450
313.25	4459
313.5	4471
313.75	4426
314	4478
314.25	4506
314.5	4522
314.75	4529
315	4580
315.25	4562
315.5	4572
315.75	4601
316	4608
316.25	4547
316.5	4560
316.75	4556
317	4603
317.25	4611
317.5	4570
317.75	4600
318	4575
318.25	4576
318.5	4568
318.75	4600
319	4553
319.25	4544
319.5	4555
319.75	4539
320	4526
320.25	4572
320.5	4528
320.75	4503
321	4508
321.25	4503
321.5	4512
321.75	4456
322	4478
322.25	4495
322.5	4485
322.75	4466
323	4435
323.25	4415
323.5	4395
323.75	4413
324	4380
324.25	4431
324.5	4428
324.75	4382
325	4455
325.25	4416
325.5	4371
325.75	4367
326	4418
326.25	4445
326.5	4450
326.75	4406
327	4409
327.25	4418
327.5	4424
327.75	4471
328	4475
328.25	4490
328.5	4458
328.75	4505
329	4511
329.25	4531
329.5	4551
329.75	4588
330	4612
330.25	4595
330.5	4594
330.75	4615
331	4622
331.25	4676
331.5	4685
331.75	4690
332	4712
332.25	4752
332.5	4775
332.75	4782
333	4798
333.25	4817
333.5	4894
333.75	4920
334	4893
334.25	4928
334.5	4990
334.75	4947
335	4934
335.25	4994
335.5	5021
335.75	5086
336	5098
336.25	5080
336.5	5094
336.75	5102
337	5150
337.25	5170
337.5	5170
337.75	5222
338	5278
338.25	5246
338.5	5315
338.75	5241
339	5372
339.25	5335
339.5	5377
339.75	5381
340	5422
340.25	5411
340.5	5437
340.75	5507
341	5461
341.25	5529
341.5	5507
341.75	5544
342	5560
342.25	5550
342.5	5556
342.75	5543
343	5597
343.25	5604
343.5	5583
343.75	5599
344	5698
344.25	5656
344.5	5667
344.75	5661
345	5671
345.25	5647
345.5	5650
345.75	5654
346	5691
346.25	5689
346.5	5633
346.75	5643
347	5689
347.25	5673
347.5	5700
347.75	5674
348	5666
348.25	5666
348.5	5692
348.75	5648
349	5677
349.25	5626
349.5	5641
349.75	5652
350	5572
350.25	5625
350.5	5613
350.75	5581
351	5577
351.25	5566
351.5	5533
351.75	5575
352	5506
352.25	5473
352.5	5480
352.75	5465
353	5435
353.25	5388
353.5	5399
353.75	5354
354	5335
354.25	5325
354.5	5316
354.75	5277
355	5238
355.25	5221
355.5	5198
355.75	5143
356	5136
356.25	5092
356.5	5032
356.75	5029
357	4995
357.25	4983
357.5	4945
357.75	4887
358	4849
358.25	4802
358.5	4814
358.75	4760
359	4751
359.25	4698
359.5	4648
359.75	4602
360	4556
360.25	4545
360.5	4487
360.75	4453
361	4405
361.25	4363
361.5	4327
361.75	4291
362	4251
362.25	4169
362.5	4150
362.75	4088
363	4053
363.25	4026
363.5	3965
363.75	3923
364	3870
364.25	3817
364.5	3778
364.75	3754
365	3691
365.25	3659
365.5	3587
365.75	3547
366	3513
366.25	3438
366.5	3384
366.75	3366
367	3297
367.25	3245
367.5	3226
367.75	3165
368	3095
368.25	3061
368.5	2993
368.75	2956
369	2919
369.25	2891
369.5	2841
369.75	2781
370	2727
370.25	2661
370.5	2622
370.75	2575
371	2523
371.25	2497
371.5	2447
371.75	2397
372	2367
372.25	2306
372.5	2273
372.75	2236
373	2189
373.25	2172
373.5	2122
373.75	2086
374	2051
374.25	2012
374.5	1984
374.75	1981
375	1958
375.25	1894
375.5	1860
375.75	1819
376	1747
376.25	1714
376.5	1659
376.75	1620
377	1604
377.25	1565
377.5	1541
377.75	1476
378	1456
378.25	1423
378.5	1373
378.75	1358
379	1317
379.25	1287
379.5	1267
379.75	1240
380	1204
380.25	1181
380.5	1149
380.75	1128
381	1100
381.25	1068
381.5	1028
381.75	1002
382	975
382.25	948
382.5	946
382.75	911
383	882
383.25	858
383.5	835
383.75	821
384	787
384.25	770
384.5	761
384.75	725
385	712
385.25	681
385.5	671
385.75	667
386	640
386.25	607
386.5	595
386.75	599
387	570
387.25	552
387.5	529
387.75	514
388	496
388.25	498
388.5	479
388.75	462
389	440
389.25	441
389.5	425
389.75	410
390	407
390.25	388
390.5	369
390.75	360
391	356
391.25	346
391.5	335
391.75	327
392	316
392.25	296
392.5	285
392.75	287
393	279
393.25	265
393.5	250
393.75	243
394	250
394.25	245
394.5	229
394.75	216
395	220
395.25	203
395.5	191
395.75	181
396	188
396.25	181
396.5	166
396.75	155
397	158
397.25	161
397.5	148
397.75	135
398	136
398.25	130
398.5	134
398.75	122
399	134
399.25	120
399.5	112
399.75	105
400	114
400.25	105
400.5	94
400.75	99
401	94
401.25	88
401.5	86
401.75	83
402	85
402.25	87
402.5	73
402.75	76
403	69
403.25	75
403.5	69
403.75	58
404	76
404.25	56
404.5	53
404.75	66
405	57
405.25	48
405.5	49
405.75	48
406	46
406.25	47
406.5	50
406.75	40
407	48
407.25	48
407.5	39
407.75	45
408	38
408.25	28
408.5	39
408.75	34
409	26
409.25	19
409.5	19
409.75	25
410	29
410.25	20
410.5	31
410.75	20
411	22
411.25	19
411.5	13
411.75	15
412	15
412.25	21
412.5	23
412.75	19
413	19
413.25	14
413.5	12
413.75	21
414	17
414.25	11
414.5	11
414.75	6
415	13
415.25	16
415.5	14
415.75	10
416	23
416.25	7
416.5	8
416.75	8
417	11
417.25	12
417.5	4
417.75	11
418	0
418.25	7
418.5	3
418.75	12
419	15
419.25	8
419.5	7
419.75	2
420	6
"""

_qs_abs_df = pd.read_csv(io.StringIO(_QUININE_SULFATE_ABS_TXT), sep="\t", header=None, names=["wavelength", "epsilon"])
QUININE_SULFATE_ABS_EXAMPLE = {
    "option_label": "Usar dato real (sulfato de quinina, PhotochemCAD)",
    "wl": _qs_abs_df["wavelength"].to_numpy(dtype=float),
    "epsilon": _qs_abs_df["epsilon"].to_numpy(dtype=float),
    "citation": (
        "Coeficiente de extincion molar real de sulfato de quinina en H\u2082SO\u2084 0.5 M "
        "(Irvin \u0026 Irvin, 1948; \u03b5=5700 M\u207b\u00b9cm\u207b\u00b9 a 347.5 nm). "
        "Digitalizado por PhotochemCAD 2.1a (Du et al. 1998; Dixon et al. 2005), "
        "reprocesado por S. Prahl, Oregon Medical Laser Center. Rango incluido: 255.5-420 nm."
    ),
    "url": "https://omlc.org/spectra/PhotochemCAD/html/081.html",
}


def area_input(label, key_prefix, default_value=1.0, min_value=0.0, help_text=None, embedded_example=None):
    """Campo de 'area integrada' reutilizable: manual, calculada subiendo un
    espectro propio, o (si se provee embedded_example) usando un dataset de
    referencia real ya incluido en la app. Devuelve el area a usar en el calculo."""
    with st.container(border=True):
        st.markdown(f"\U0001F4C8 **{label}**")
        if help_text:
            st.caption(help_text)

        options = ["Manual", "Subir espectro"]
        if embedded_example:
            options.append(embedded_example["option_label"])
        mode = st.segmented_control(
            "Origen del dato", options, default=options[0], required=True,
            key=f"{key_prefix}_mode", label_visibility="collapsed",
        )

        if mode == "Manual":
            return st.number_input(label, min_value=min_value, value=default_value,
                                   format="%.6f", key=f"{key_prefix}_manual", label_visibility="collapsed")

        if embedded_example and mode == embedded_example["option_label"]:
            st.info(f"\U0001F4DA {embedded_example['citation']}")
            st.link_button("\U0001F517 Ver la fuente original", embedded_example["url"])
            try:
                return _render_spectrum_result(embedded_example["wl"], embedded_example["intensity"])
            except Exception as e:
                st.error(f"No se pudo calcular el area del dataset de referencia: {e}")
                return default_value

        f = st.file_uploader("Espectro (CSV o XLSX)", type=["csv", "xlsx"],
                             key=f"{key_prefix}_file", label_visibility="collapsed")
        if f is None:
            st.caption("Sube un archivo con columnas de longitud de onda e intensidad.")
            return default_value

        try:
            raw = f.getvalue()
            if f.name.lower().endswith(".xlsx"):
                sheets = list_excel_sheets(raw)
                sheet = st.selectbox("Hoja", options=sheets, key=f"{key_prefix}_sheet")
                df = load_excel_sheet(raw, sheet)
            else:
                df = load_csv_df(raw)

            wl_guess, int_guess = guess_columns(df)
            cols = list(df.columns)
            with st.expander("\u2699\ufe0f Ajustar columnas, rango y normalizacion", expanded=False):
                c1, c2, c3 = st.columns(3)
                with c1:
                    wl_col = st.selectbox("Columna wavelength", options=cols,
                                          index=cols.index(wl_guess) if wl_guess in cols else 0,
                                          key=f"{key_prefix}_wlcol")
                with c2:
                    int_col = st.selectbox("Columna intensidad", options=cols,
                                           index=cols.index(int_guess) if int_guess in cols else min(1, len(cols) - 1),
                                           key=f"{key_prefix}_intcol")
                with c3:
                    normalize = st.selectbox("Normalizacion", options=["None", "Max = 1", "Area = 1"],
                                             key=f"{key_prefix}_norm")
                c4, c5 = st.columns(2)
                with c4:
                    wl_min = st.number_input("Min nm", min_value=100, max_value=10000, value=200, key=f"{key_prefix}_wlmin")
                with c5:
                    wl_max = st.number_input("Max nm", min_value=100, max_value=10000, value=900, key=f"{key_prefix}_wlmax")

            wl, intensity = filter_and_normalize(df[wl_col], df[int_col], wl_min, wl_max, normalize)
            return _render_spectrum_result(wl, intensity)
        except Exception as e:
            st.error(f"No se pudo calcular el area a partir del espectro: {e}")
            return default_value


def absorbance_input(label, key_prefix, default_value=0.05, excitation_wl_default=350.0, help_text=None, embedded_example=None):
    """Campo de 'absorbancia' reutilizable: manual, leida de un espectro de
    absorcion/reflectancia propio, o (si se provee embedded_example) calculada
    a partir de un coeficiente de extincion molar real ya incluido en la app
    (requiere que el usuario indique la concentracion y el paso optico usados
    en su propia medicion). Devuelve el valor a usar en el calculo."""
    with st.container(border=True):
        st.markdown(f"\U0001F52C **{label}**")
        if help_text:
            st.caption(help_text)

        options = ["Manual", "Subir espectro"]
        if embedded_example:
            options.append(embedded_example["option_label"])
        mode = st.segmented_control(
            "Origen del dato", options, default=options[0], required=True,
            key=f"{key_prefix}_mode", label_visibility="collapsed",
        )

        if mode == "Manual":
            return st.number_input(label, min_value=0.000001, value=default_value,
                                   format="%.6f", key=f"{key_prefix}_manual", label_visibility="collapsed")

        if embedded_example and mode == embedded_example["option_label"]:
            st.info(f"\U0001F4DA {embedded_example['citation']}")
            st.link_button("\U0001F517 Ver la fuente original", embedded_example["url"])
            st.caption(
                "El coeficiente de extincion molar (\u03b5) es un dato real y publicado, pero la "
                "absorbancia final depende de la concentracion y el paso optico de tu propia "
                "medicion (A = \u03b5\u00b7c\u00b7l), que no vienen publicados junto al espectro. "
                "Ingresalos abajo para completar el calculo."
            )
            try:
                wl = embedded_example["wl"]
                eps = embedded_example["epsilon"]
                exc_wl = st.number_input(
                    "Longitud de onda de excitacion (nm)",
                    min_value=float(wl.min()), max_value=float(wl.max()),
                    value=float(np.clip(excitation_wl_default, wl.min(), wl.max())),
                    format="%.1f", key=f"{key_prefix}_exwl_embed",
                )
                c1, c2 = st.columns(2)
                with c1:
                    conc = st.number_input("Concentracion (mol/L)", min_value=1e-9, value=1e-5,
                                           format="%.3e", key=f"{key_prefix}_conc_embed")
                with c2:
                    path_len = st.number_input("Paso optico (cm)", min_value=0.01, value=1.0,
                                               format="%.3f", key=f"{key_prefix}_path_embed")

                epsilon_at_wl = float(np.interp(exc_wl, wl, eps))
                absorbance = epsilon_at_wl * conc * path_len

                fig, ax = plt.subplots(figsize=(4, 2.0))
                ax.plot(wl, eps, color="#d62728", linewidth=1.6)
                ax.fill_between(wl, eps, color="#d62728", alpha=0.10)
                ax.axvline(exc_wl, color="#6b7280", linestyle="--", linewidth=1)
                ax.set_xlabel("Longitud de onda (nm)", fontsize=8)
                ax.set_ylabel("\u03b5 (M\u207b\u00b9cm\u207b\u00b9)", fontsize=8)
                ax.tick_params(labelsize=7)
                ax.grid(alpha=0.2)
                show_and_close(fig)

                m1, m2, m3 = st.columns(3)
                m1.metric("\u03b5 en exc.", f"{epsilon_at_wl:.4g}")
                m2.metric("A = \u03b5\u00b7c\u00b7l", f"{absorbance:.4g}")
                m3.metric("\u2265 0.1", "\u26a0\ufe0f Si" if absorbance >= 0.1 else "\u2705 No")
                if absorbance >= 0.1:
                    st.warning(
                        "Esta absorbancia es \u2265 0.1: reduce la concentracion para evitar "
                        "efectos de filtro interno en el rendimiento cuantico relativo."
                    )
                return absorbance
            except Exception as e:
                st.error(f"No se pudo calcular la absorbancia a partir de \u03b5: {e}")
                return default_value

        f = st.file_uploader("Espectro de absorcion o reflectancia (CSV o XLSX)", type=["csv", "xlsx"],
                             key=f"{key_prefix}_file", label_visibility="collapsed")
        if f is None:
            st.caption("Sube un archivo con columnas de longitud de onda y absorbancia/reflectancia.")
            return default_value

        try:
            raw = f.getvalue()
            if f.name.lower().endswith(".xlsx"):
                sheets = list_excel_sheets(raw)
                sheet = st.selectbox("Hoja", options=sheets, key=f"{key_prefix}_sheet")
                df = load_excel_sheet(raw, sheet)
            else:
                df = load_csv_df(raw)

            wl_guess, val_guess = guess_columns(df)
            cols = list(df.columns)
            with st.expander("\u2699\ufe0f Ajustar columnas y tipo de espectro", expanded=False):
                c1, c2, c3 = st.columns(3)
                with c1:
                    wl_col = st.selectbox("Columna wavelength", options=cols,
                                          index=cols.index(wl_guess) if wl_guess in cols else 0,
                                          key=f"{key_prefix}_wlcol")
                with c2:
                    val_col = st.selectbox("Columna de valor", options=cols,
                                           index=cols.index(val_guess) if val_guess in cols else min(1, len(cols) - 1),
                                           key=f"{key_prefix}_valcol")
                with c3:
                    spectrum_kind = st.selectbox(
                        "Tipo de espectro",
                        options=["Absorbancia", "Reflectancia (%)", "Reflectancia (fraccion 0-1)"],
                        key=f"{key_prefix}_kind",
                    )
                exc_wl = st.number_input("Longitud de onda de excitacion (nm)", min_value=100.0, max_value=2000.0,
                                         value=excitation_wl_default, format="%.1f", key=f"{key_prefix}_excwl")

            wl = to_numeric_series(df[wl_col]).to_numpy(dtype=float)
            val = to_numeric_series(df[val_col]).to_numpy(dtype=float)
            mask = np.isfinite(wl) & np.isfinite(val)
            wl, val = wl[mask], val[mask]
            order = np.argsort(wl)
            wl, val = wl[order], val[order]
            if len(wl) < 2:
                raise ValueError("El espectro no tiene suficientes puntos numericos validos.")

            raw_value = float(np.interp(exc_wl, wl, val))

            fig, ax = plt.subplots(figsize=(4, 2.0))
            ax.plot(wl, val, color="#d62728", linewidth=1.6)
            ax.fill_between(wl, val, color="#d62728", alpha=0.10)
            ax.axvline(exc_wl, color="#6b7280", linestyle="--", linewidth=1)
            ax.set_xlabel("Longitud de onda (nm)", fontsize=8)
            ax.set_ylabel(spectrum_kind, fontsize=8)
            ax.tick_params(labelsize=7)
            ax.grid(alpha=0.2)
            show_and_close(fig)

            if spectrum_kind == "Absorbancia":
                absorbance = raw_value
                st.metric(f"Absorbancia a {exc_wl:.1f} nm", f"{absorbance:.4g}")
            else:
                R = raw_value / 100.0 if spectrum_kind == "Reflectancia (%)" else raw_value
                R = min(max(R, 1e-6), 0.999999)
                absorbance = ((1 - R) ** 2) / (2 * R)
                m1, m2 = st.columns(2)
                m1.metric(f"Reflectancia a {exc_wl:.1f} nm", f"{R:.4f}")
                m2.metric("F(R) Kubelka-Munk", f"{absorbance:.4g}")
                st.caption(
                    "F(R) = (1-R)\u00b2 / (2R), usada en lugar de la absorbancia para muestras solidas "
                    "en la ecuacion de rendimiento cuantico relativo."
                )
            return absorbance
        except Exception as e:
            st.error(f"No se pudo obtener la absorbancia a partir del espectro: {e}")
            return default_value


def _docx_report_header(doc, title, timestamp, intro_text):
    heading = doc.add_heading(title, level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x7F, 0x1D, 0x1D)
    doc.add_paragraph(f"Generado: {timestamp}")
    doc.add_paragraph(intro_text)


def _docx_add_step(doc, step_number, step_title, concept, lines):
    doc.add_heading(f"Paso {step_number}. {step_title}", level=2)
    if concept:
        doc.add_paragraph(concept)
    for line in lines:
        p = doc.add_paragraph()
        p.add_run(line).italic = True


def _docx_add_result(doc, symbol, value, warn_if_over_one=True):
    doc.add_heading("Resultado", level=2)
    result_p = doc.add_paragraph()
    result_run = result_p.add_run(f"{symbol} = {value:.4f}  ({value * 100:.2f} %)")
    result_run.bold = True
    result_run.font.size = Pt(14)
    if warn_if_over_one and value > 1:
        warn = doc.add_paragraph()
        warn.add_run(
            "Advertencia: el resultado es mayor que 1. Revisa areas, absorbancias/intensidades, "
            "referencia o correcciones experimentales."
        ).italic = True


def _pdf_header_story(styles, title, timestamp, intro_text):
    title_style = styles["Title"]
    title_style.textColor = rl_colors.HexColor("#7F1D1D")
    return [
        Paragraph(title, title_style),
        Spacer(1, 6),
        Paragraph(f"Generado: {timestamp}", styles["Normal"]),
        Spacer(1, 10),
        Paragraph(intro_text, styles["Normal"]),
        Spacer(1, 14),
    ]


def _pdf_step_story(styles, step_number, step_title, concept, lines):
    story = [Paragraph(f"<b>Paso {step_number}. {step_title}</b>", styles["Heading2"])]
    if concept:
        story.append(Paragraph(concept, styles["Normal"]))
        story.append(Spacer(1, 4))
    for line in lines:
        story.append(Paragraph(f"<i>{line}</i>", styles["Normal"]))
        story.append(Spacer(1, 4))
    story.append(Spacer(1, 10))
    return story


def _pdf_result_story(styles, symbol, value, warn_if_over_one=True):
    story = [
        Paragraph("<b>Resultado</b>", styles["Heading2"]),
        Paragraph(
            f"<font size=14><b>{symbol} = {value:.4f}  ({value * 100:.2f} %)</b></font>",
            styles["Normal"],
        ),
    ]
    if warn_if_over_one and value > 1:
        story.append(Spacer(1, 10))
        story.append(Paragraph(
            "<i>Advertencia: el resultado es mayor que 1. Revisa areas, absorbancias/intensidades, "
            "referencia o correcciones experimentales.</i>",
            styles["Normal"],
        ))
    return story


# ---------- Rendimiento cuantico RELATIVO ----------

def _relative_qy_steps(data):
    """Calcula las razones intermedias y arma el texto de cada paso,
    compartido entre el exportador a Word y a PDF."""
    ratio_area = data["sample_area"] / data["ref_area"]
    ratio_abs = data["ref_abs"] / data["sample_abs"]
    ratio_n2 = (data["sample_n"] ** 2) / (data["ref_n"] ** 2)
    phi = data["ref_phi"] * ratio_area * ratio_abs * ratio_n2

    steps = [
        (1, "Formula general del metodo relativo",
         "El metodo relativo compara la muestra con una sustancia de referencia cuyo rendimiento "
         "cuantico ya se conoce de la literatura, midiendo ambas bajo la misma longitud de onda de "
         "excitacion y, si es posible, con absorbancias bajas y similares.",
         [
            "\u03a6x = \u03a6ref \u00d7 (Ix / Iref) \u00d7 (Aref / Ax) \u00d7 (nx\u00b2 / nref\u00b2)",
        ]),
        (2, "Razon de areas de emision integradas (Ix / Iref)",
         "Esta razon compara cuanta luz emite la muestra frente a la referencia. Un valor mayor que 1 "
         "significa que la muestra emite mas area de fluorescencia que la referencia, en igualdad de condiciones.",
         [
            f"Ix / Iref = {data['sample_area']:.6g} / {data['ref_area']:.6g} = {ratio_area:.6g}",
        ]),
        (3, "Razon de absorbancias a la longitud de onda de excitacion (Aref / Ax)",
         "Esta razon corrige el hecho de que la muestra y la referencia no absorben exactamente la "
         "misma cantidad de luz de excitacion; sin esta correccion, la comparacion del paso anterior no seria justa.",
         [
            f"Aref / Ax = {data['ref_abs']:.6g} / {data['sample_abs']:.6g} = {ratio_abs:.6g}",
        ]),
        (4, "Razon de indices de refraccion al cuadrado (nx\u00b2 / nref\u00b2)",
         "Corrige por el solvente: la luz se refracta distinto en cada solvente, lo que afecta cuanta "
         "emision llega efectivamente al detector. Si ambas muestras usan el mismo solvente, esta razon vale 1.",
         [
            f"nx\u00b2 / nref\u00b2 = {data['sample_n']:.6g}\u00b2 / {data['ref_n']:.6g}\u00b2 "
            f"= {data['sample_n'] ** 2:.6g} / {data['ref_n'] ** 2:.6g} = {ratio_n2:.6g}",
        ]),
        (5, "Sustituir todas las razones en la formula",
         "Con las tres correcciones ya calculadas, se combinan multiplicando el rendimiento cuantico "
         "conocido de la referencia por cada una de ellas.",
         [
            f"\u03a6x = {data['ref_phi']:.6g} \u00d7 {ratio_area:.6g} \u00d7 {ratio_abs:.6g} \u00d7 {ratio_n2:.6g}",
        ]),
        (6, "Multiplicar y obtener el resultado",
         "El numero final es el rendimiento cuantico relativo de la muestra: la fraccion de fotones "
         "absorbidos que la muestra reemite como fluorescencia.",
         [
            f"\u03a6x = {phi:.6g}",
        ]),
    ]
    return phi, steps


def build_relative_qy_docx(data: dict) -> bytes:
    """Genera un reporte .docx con el desarrollo matematico paso a paso
    del rendimiento cuantico relativo."""
    phi, steps = _relative_qy_steps(data)
    doc = Document()
    _docx_report_header(
        doc, "Rendimiento cuantico relativo", data["timestamp"],
        "Metodo relativo: compara el area de emision integrada de la muestra con la de "
        "una referencia de rendimiento cuantico conocido, a la misma longitud de onda de excitacion.",
    )

    doc.add_heading("Datos de entrada", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Parametro"
    hdr[1].text = "Muestra"
    hdr[2].text = "Referencia"
    rows = [
        ("Area integrada de emision", f"{data['sample_area']:.6g}", f"{data['ref_area']:.6g}"),
        ("Absorbancia (exc.)", f"{data['sample_abs']:.6g}", f"{data['ref_abs']:.6g}"),
        ("Indice de refraccion", f"{data['sample_n']:.6g}", f"{data['ref_n']:.6g}"),
        ("Rendimiento cuantico (Phi)", "-", f"{data['ref_phi']:.6g}"),
    ]
    for label, sample_val, ref_val in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = sample_val
        cells[2].text = ref_val

    doc.add_heading("Desarrollo matematico", level=2)
    for step_number, step_title, concept, lines in steps:
        _docx_add_step(doc, step_number, step_title, concept, lines)

    _docx_add_result(doc, "\u03a6x", phi)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def build_relative_qy_pdf(data: dict) -> bytes:
    """Genera un reporte .pdf con el desarrollo matematico paso a paso
    del rendimiento cuantico relativo."""
    phi, steps = _relative_qy_steps(data)
    bio = io.BytesIO()
    doc = SimpleDocTemplate(bio, pagesize=letter, topMargin=0.8 * rl_inch, bottomMargin=0.8 * rl_inch)
    styles = getSampleStyleSheet()

    story = _pdf_header_story(
        styles, "Rendimiento cuantico relativo", data["timestamp"],
        "Metodo relativo: compara el area de emision integrada de la muestra con la de "
        "una referencia de rendimiento cuantico conocido, a la misma longitud de onda de excitacion.",
    )

    story.append(Paragraph("<b>Datos de entrada</b>", styles["Heading2"]))
    table_data = [
        ["Parametro", "Muestra", "Referencia"],
        ["Area integrada de emision", f"{data['sample_area']:.6g}", f"{data['ref_area']:.6g}"],
        ["Absorbancia (exc.)", f"{data['sample_abs']:.6g}", f"{data['ref_abs']:.6g}"],
        ["Indice de refraccion", f"{data['sample_n']:.6g}", f"{data['ref_n']:.6g}"],
        ["Rendimiento cuantico (Phi)", "-", f"{data['ref_phi']:.6g}"],
    ]
    tbl = Table(table_data, colWidths=[2.4 * rl_inch, 1.6 * rl_inch, 1.6 * rl_inch])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), rl_colors.HexColor("#FEE2E2")),
        ("TEXTCOLOR", (0, 0), (-1, 0), rl_colors.HexColor("#7F1D1D")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.5, rl_colors.HexColor("#E5E7EB")),
        ("ALIGN", (1, 0), (-1, -1), "CENTER"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [rl_colors.white, rl_colors.HexColor("#F9FAFB")]),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 16))

    story.append(Paragraph("<b>Desarrollo matematico</b>", styles["Heading2"]))
    story.append(Spacer(1, 6))
    for step_number, step_title, concept, lines in steps:
        story.extend(_pdf_step_story(styles, step_number, step_title, concept, lines))

    story.extend(_pdf_result_story(styles, "\u03a6x", phi))

    doc.build(story)
    bio.seek(0)
    return bio.getvalue()


# ---------- Rendimiento cuantico ABSOLUTO (metodo de esfera integradora) ----------

def _absolute_qy_steps(data):
    """Metodo de de Mello, Wittmann y Friend (1997) para rendimiento cuantico
    absoluto por esfera integradora. Calcula pasos intermedios y arma el
    texto de cada paso, compartido entre el exportador a Word y a PDF."""
    La, Lc = data["L_a"], data["L_c"]
    Ea, Ec = data["E_a"], data["E_c"]

    lc_la_ratio = Lc / La
    A = 1 - lc_la_ratio
    one_minus_A_times_Ea = (1 - A) * Ea
    numerator = Ec - one_minus_A_times_Ea
    denominator = A * La
    phi_abs = numerator / denominator if denominator != 0 else float("nan")

    steps = [
        (1, "Fraccion de luz absorbida (A)",
         "Antes de calcular el rendimiento cuantico hay que saber que fraccion de la luz de excitacion "
         "realmente absorbio la muestra. Se obtiene comparando cuanta luz de excitacion se detecta con "
         "la muestra puesta (Lc) frente a sin ella, solo el blanco/referencia (La).",
         [
            "A = 1 \u2212 (Lc / La)",
            f"Lc / La = {Lc:.6g} / {La:.6g} = {lc_la_ratio:.6g}",
            f"A = 1 \u2212 {lc_la_ratio:.6g} = {A:.6g}",
        ]),
        (2, "Formula general del rendimiento cuantico absoluto",
         "A diferencia del metodo relativo, este metodo no necesita una sustancia de referencia externa: "
         "toda la informacion sale de cuatro mediciones hechas con la misma esfera integradora.",
         [
            "\u03a6abs = (Ec \u2212 (1 \u2212 A)\u00b7Ea) / (A\u00b7La)",
        ]),
        (3, "Termino de emision indirecta corregido: (1 \u2212 A)\u00b7Ea",
         "Este termino corrige la emision que ocurriria por excitacion indirecta (luz dispersada dentro "
         "de la esfera que tambien excita la muestra), para no contarla dos veces en el resultado final.",
         [
            f"(1 \u2212 A) = 1 \u2212 {A:.6g} = {1 - A:.6g}",
            f"(1 \u2212 A)\u00b7Ea = {1 - A:.6g} \u00d7 {Ea:.6g} = {one_minus_A_times_Ea:.6g}",
        ]),
        (4, "Numerador: Ec \u2212 (1 \u2212 A)\u00b7Ea",
         "El numerador representa la emision de fluorescencia que puede atribuirse unicamente a la "
         "excitacion directa de la muestra, ya sin la contribucion indirecta del paso anterior.",
         [
            f"Numerador = {Ec:.6g} \u2212 {one_minus_A_times_Ea:.6g} = {numerator:.6g}",
        ]),
        (5, "Denominador: A\u00b7La",
         "El denominador representa el total de fotones de excitacion que la muestra efectivamente absorbio.",
         [
            f"Denominador = {A:.6g} \u00d7 {La:.6g} = {denominator:.6g}",
        ]),
        (6, "Dividir numerador entre denominador",
         "El cociente entre fotones emitidos (numerador) y fotones absorbidos (denominador) es, por "
         "definicion, el rendimiento cuantico absoluto de la muestra.",
         [
            f"\u03a6abs = {numerator:.6g} / {denominator:.6g} = {phi_abs:.6g}",
        ]),
    ]
    return phi_abs, A, steps


def build_absolute_qy_docx(data: dict) -> bytes:
    """Genera un reporte .docx con el desarrollo matematico paso a paso
    del rendimiento cuantico absoluto (metodo de esfera integradora)."""
    phi_abs, A, steps = _absolute_qy_steps(data)
    doc = Document()
    _docx_report_header(
        doc, "Rendimiento cuantico absoluto", data["timestamp"],
        "Metodo absoluto (de Mello, Wittmann y Friend, 1997) con esfera integradora: usa las "
        "areas de excitacion dispersada y de emision, medidas con y sin excitacion directa de "
        "la muestra, para obtener el rendimiento cuantico sin necesitar una referencia externa."
    )

    doc.add_heading("Datos de entrada", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Parametro"
    hdr[1].text = "Valor (area integrada)"
    rows = [
        ("La \u2014 excitacion, referencia/blanco", f"{data['L_a']:.6g}"),
        ("Lc \u2014 excitacion, con la muestra", f"{data['L_c']:.6g}"),
        ("Ea \u2014 emision, excitacion indirecta", f"{data['E_a']:.6g}"),
        ("Ec \u2014 emision, excitacion directa", f"{data['E_c']:.6g}"),
    ]
    for label, val in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = val

    doc.add_heading("Desarrollo matematico", level=2)
    for step_number, step_title, concept, lines in steps:
        _docx_add_step(doc, step_number, step_title, concept, lines)

    _docx_add_result(doc, "\u03a6abs", phi_abs)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def build_absolute_qy_pdf(data: dict) -> bytes:
    """Genera un reporte .pdf con el desarrollo matematico paso a paso
    del rendimiento cuantico absoluto (metodo de esfera integradora)."""
    phi_abs, A, steps = _absolute_qy_steps(data)
    bio = io.BytesIO()
    doc = SimpleDocTemplate(bio, pagesize=letter, topMargin=0.8 * rl_inch, bottomMargin=0.8 * rl_inch)
    styles = getSampleStyleSheet()

    story = _pdf_header_story(
        styles, "Rendimiento cuantico absoluto", data["timestamp"],
        "Metodo absoluto (de Mello, Wittmann y Friend, 1997) con esfera integradora: usa las "
        "areas de excitacion dispersada y de emision, medidas con y sin excitacion directa de "
        "la muestra, para obtener el rendimiento cuantico sin necesitar una referencia externa."
    )

    story.append(Paragraph("<b>Datos de entrada</b>", styles["Heading2"]))
    table_data = [
        ["Parametro", "Valor (area integrada)"],
        ["La \u2014 excitacion, referencia/blanco", f"{data['L_a']:.6g}"],
        ["Lc \u2014 excitacion, con la muestra", f"{data['L_c']:.6g}"],
        ["Ea \u2014 emision, excitacion indirecta", f"{data['E_a']:.6g}"],
        ["Ec \u2014 emision, excitacion directa", f"{data['E_c']:.6g}"],
    ]
    tbl = Table(table_data, colWidths=[3.2 * rl_inch, 2.4 * rl_inch])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), rl_colors.HexColor("#FEE2E2")),
        ("TEXTCOLOR", (0, 0), (-1, 0), rl_colors.HexColor("#7F1D1D")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.5, rl_colors.HexColor("#E5E7EB")),
        ("ALIGN", (1, 0), (-1, -1), "CENTER"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [rl_colors.white, rl_colors.HexColor("#F9FAFB")]),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 16))

    story.append(Paragraph("<b>Desarrollo matematico</b>", styles["Heading2"]))
    story.append(Spacer(1, 6))
    for step_number, step_title, concept, lines in steps:
        story.extend(_pdf_step_story(styles, step_number, step_title, concept, lines))

    story.extend(_pdf_result_story(styles, "\u03a6abs", phi_abs))

    doc.build(story)
    bio.seek(0)
    return bio.getvalue()


# ============================================================
# Navegacion global (sidebar) - vacia en Inicio, con enlaces en el resto
# ============================================================
PAGES = ["Inicio", "Analisis CIE 1931", "Visor de espectros", "Rendimiento cuantico", "Aprender"]
PAGE_ICONS = {
    "Inicio": "\U0001F3E0",
    "Analisis CIE 1931": "\U0001F3A8",
    "Visor de espectros": "\U0001F4C8",
    "Rendimiento cuantico": "\u269B\uFE0F",
    "Aprender": "\U0001F4DA",
}

if st.session_state["active_page"] != "Inicio":
    st.sidebar.markdown(
        """
        <style>
        div[data-testid="stSidebar"] div[data-testid="stButton"] > button {
            text-align: left;
            border-radius: 8px;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )
    st.sidebar.header("Navegacion")
    for _page_name in PAGES:
        _icon = PAGE_ICONS.get(_page_name, "")
        if st.sidebar.button(f"{_icon}  {_page_name}", use_container_width=True, key=f"nav_{_page_name}"):
            go_to_page(_page_name)
    st.sidebar.markdown("---")
else:
    st.sidebar.markdown(
        """
        <style>
        div[class*="st-key-toolbox_"] {
            border: 1px solid #e5e7eb;
            border-radius: 10px;
            padding: .55rem .7rem;
            margin-bottom: .6rem;
            background-color: #fafafa;
            transition: border-color .2s ease, background-color .2s ease, box-shadow .2s ease;
        }
        div[class*="st-key-toolbox_"]:hover {
            border-color: #fca5a5;
            background-color: #fff5f5;
            box-shadow: 0 4px 14px rgba(0, 0, 0, 0.07);
        }
        div[class*="st-key-toolbox_"] .tool-label {
            font-weight: 600;
            font-size: 1.1rem;
            color: #31333F;
            transition: color .2s ease;
        }
        div[class*="st-key-toolbox_"]:hover .tool-label {
            color: #b91c1c;
        }
        div[class*="st-key-toolbox_"] div[data-testid="stHorizontalBlock"] {
            margin-top: .5rem;
        }
        div[class*="st-key-toolbox_"] div[data-testid="stButton"] > button {
            border-radius: 7px;
            font-size: 1rem;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )
    st.sidebar.header("Herramientas")

    _tools = [
        ("Analisis CIE 1931", "cie", "CIE 1931"),
        ("Visor de espectros", "visor", "Visor de espectros"),
        ("Rendimiento cuantico", "rendimiento", "Rendimiento cuantico"),
    ]
    for _tool_name, _tool_key, _learn_topic in _tools:
        _icon = PAGE_ICONS.get(_tool_name, "")
        try:
            _tool_box = st.sidebar.container(key=f"toolbox_{_tool_key}")
        except TypeError:
            # Streamlit antiguo sin soporte para container(key=...): pierde el
            # hover-reveal pero sigue funcionando como lista simple.
            _tool_box = st.sidebar.container()
        with _tool_box:
            st.markdown(f'<div class="tool-label">{_icon}&nbsp;&nbsp;{_tool_name}</div>', unsafe_allow_html=True)
            _b1, _b2 = st.columns(2)
            with _b1:
                if st.button("Calcular", key=f"{_tool_key}_calc", use_container_width=True):
                    go_to_page(_tool_name)
            with _b2:
                if st.button("Aprender", key=f"{_tool_key}_learn", use_container_width=True):
                    st.session_state["learn_topic"] = _learn_topic
                    go_to_page("Aprender")

# ============================================================
# Pagina: Inicio
# ============================================================
if st.session_state["active_page"] == "Inicio":
    st.markdown(
        """
        <style>
        .main .block-container {
            padding-top: 2rem;
            padding-bottom: 1rem;
        }
        .main .block-container h1 { margin-bottom: 0rem; }
        .main .block-container h3, .main .block-container h4 { margin-top: 0.25rem; margin-bottom: 0.5rem; }
        </style>
        """,
        unsafe_allow_html=True,
    )
    st.title("SpectraLab Toolkit")
    st.caption("Conjunto de herramientas para visualizar, comparar y analizar datos espectroscopicos.")

    st.write("")
    box_left, box_right = st.columns(2, gap="large")

    with box_left:
        with st.container(border=True):
            st.markdown("#### Citanos")
            st.markdown(
                "_Pendiente: agrega aqui la referencia bibliografica de tu articulo o publicacion._"
            )
            st.write("")
            st.markdown("#### Contacto")
            st.markdown("[ma.maciasl@uniandes.edu.co](mailto:ma.maciasl@uniandes.edu.co)")
            st.write("camiloserrano02@hotmail.com")

    with box_right:
        with st.container(border=True):
            st.markdown("#### Logos")
            try:
                st.image("logog.jpg", use_container_width=True)
            except Exception:
                st.info(
                    "No se encontro el archivo 'logog.jpg'. Colocalo en la misma carpeta "
                    "que app.py para que el logo aparezca aqui."
                )

    st.stop()

# ============================================================
# Pagina: Aprender
# ============================================================
if st.session_state["active_page"] == "Aprender":
    if "learn_topic" not in st.session_state:
        st.session_state["learn_topic"] = "CIE 1931"

    st.title("Aprender")
    _topics = ["CIE 1931", "Visor de espectros", "Rendimiento cuantico"]
    _default_topic = st.session_state["learn_topic"] if st.session_state["learn_topic"] in _topics else _topics[0]
    topic = st.radio("Tema", options=_topics, index=_topics.index(_default_topic),
                     horizontal=True, key="learn_topic_radio")
    st.session_state["learn_topic"] = topic
    st.markdown("---")

    if topic == "CIE 1931":
        st.markdown("### Que son las coordenadas CIE 1931")
        st.write(
            "CIE 1931 es uno de los sistemas colorimetricos mas usados para representar colores "
            "a partir de la respuesta visual humana. Fue publicado en 1931 por la Commission "
            "Internationale de l'Eclairage, tambien conocida como CIE."
        )
        st.write(
            "El diagrama CIE 1931 usa dos coordenadas, x e y, para describir la cromaticidad "
            "de una fuente luminosa. En este caso, la fuente es el espectro de emision que subes "
            "a la aplicacion."
        )
        st.write(
            "La ventaja de este sistema es que separa la informacion de color de la intensidad total. "
            "Por eso dos espectros con intensidades distintas pueden tener coordenadas x,y similares "
            "si su distribucion espectral produce una percepcion de color parecida."
        )
        st.markdown("### Para que sirve en espectros de emision")
        st.write(
            "En materiales luminiscentes, LEDs, fosforos, colorantes o complejos emisores, "
            "las coordenadas CIE permiten comparar rapidamente el color emitido por diferentes muestras."
        )
        st.markdown("### Como se calculan")
        st.write(
            "El calculo parte del espectro de emision: longitud de onda contra intensidad. "
            "Ese espectro se combina con las funciones de igualacion de color del observador "
            "estandar CIE 1931 de 2 grados."
        )
        st.latex(r"X = \int I(\lambda)\,\overline{x}(\lambda)\,d\lambda")
        st.latex(r"Y = \int I(\lambda)\,\overline{y}(\lambda)\,d\lambda")
        st.latex(r"Z = \int I(\lambda)\,\overline{z}(\lambda)\,d\lambda")
        st.write(
            "Aqui, I(lambda) es la intensidad del espectro de emision, y xbar, ybar, zbar "
            "son las funciones colorimetricas CIE. En el codigo, estas integrales se calculan "
            "numericamente con la regla trapezoidal."
        )
        st.write("Despues se normalizan los valores XYZ para obtener las coordenadas cromaticas:")
        st.latex(r"x = \frac{X}{X + Y + Z}")
        st.latex(r"y = \frac{Y}{X + Y + Z}")
        st.write(
            "Esas coordenadas x,y son las que se ubican sobre el diagrama CIE 1931. "
            "La longitud dominante y la pureza se estiman trazando una linea desde el punto blanco "
            "hasta la coordenada de la muestra y buscando su interseccion con el borde del diagrama."
        )
        st.markdown("### Ejemplo interactivo")
        st.write(
            "Sube un archivo con dos columnas: longitud de onda en nm e intensidad. "
            "La aplicacion interpola el espectro, multiplica cada punto por las funciones CIE "
            "y suma esas contribuciones para obtener X, Y y Z."
        )

        data_source = st.radio(
            "Datos para el ejemplo",
            options=["Usar ejemplo incluido", "Subir mis datos"],
            horizontal=True,
            key="cie_example_source",
        )

        example_file = None
        if data_source == "Subir mis datos":
            example_file = st.file_uploader(
                "Subir datos para el ejemplo",
                type=["csv", "xlsx"],
                key="cie_example_file"
            )

        if data_source == "Usar ejemplo incluido" or example_file is not None:
            try:
                if data_source == "Usar ejemplo incluido":
                    sample_wl = np.arange(380, 781, 5, dtype=float)
                    sample_intensity = (
                        1.00 * np.exp(-0.5 * ((sample_wl - 525.0) / 32.0) ** 2)
                        + 0.35 * np.exp(-0.5 * ((sample_wl - 610.0) / 45.0) ** 2)
                    )
                    df_example = pd.DataFrame({
                        "wavelength_nm": sample_wl,
                        "intensity": sample_intensity,
                    })
                else:
                    raw_example = example_file.getvalue()
                    if example_file.name.lower().endswith(".xlsx"):
                        sheet_names = list_excel_sheets(raw_example)
                        sheet_name = st.selectbox("Hoja de Excel", options=sheet_names, key="cie_example_sheet")
                        df_example = load_excel_sheet(raw_example, sheet_name)
                    else:
                        df_example = read_csv_flexible_bytes(raw_example)

                wl_guess, int_guess = guess_columns(df_example)
                cols = list(df_example.columns)
                c_a, c_b, c_c = st.columns(3)
                with c_a:
                    wl_col = st.selectbox(
                        "Columna de longitud de onda",
                        options=cols,
                        index=cols.index(wl_guess) if wl_guess in cols else 0,
                        key="cie_example_wl_col"
                    )
                with c_b:
                    int_col = st.selectbox(
                        "Columna de intensidad",
                        options=cols,
                        index=cols.index(int_guess) if int_guess in cols else min(1, len(cols) - 1),
                        key="cie_example_int_col"
                    )
                with c_c:
                    interval = st.selectbox("Intervalo de interpolacion (nm)", options=[1, 2, 5, 10], index=2, key="cie_example_interval")

                c_d, c_e = st.columns(2)
                with c_d:
                    wl_min = st.number_input("Longitud minima (nm)", min_value=200, max_value=10000, value=380, key="cie_example_wl_min")
                with c_e:
                    wl_max = st.number_input("Longitud maxima (nm)", min_value=200, max_value=10000, value=780, key="cie_example_wl_max")

                data = pd.DataFrame({
                    "wavelength": to_numeric_series(df_example[wl_col]),
                    "intensity": to_numeric_series(df_example[int_col]),
                }).dropna()
                data = data[(data["wavelength"] >= wl_min) & (data["wavelength"] <= wl_max)].copy()
                if data.empty:
                    raise ValueError("No hay datos dentro del rango seleccionado.")
                data = data.sort_values("wavelength").groupby("wavelength", as_index=False)["intensity"].mean()

                wl_grid = np.arange(float(wl_min), float(wl_max) + 1e-9, float(interval), dtype=float)
                intensity_grid = np.interp(
                    wl_grid,
                    data["wavelength"].values.astype(float),
                    data["intensity"].values.astype(float),
                    left=0.0,
                    right=0.0,
                )

                xbar = np.interp(wl_grid, CMF_WLS, CMF_X)
                ybar = np.interp(wl_grid, CMF_WLS, CMF_Y)
                zbar = np.interp(wl_grid, CMF_WLS, CMF_Z)
                x_product = intensity_grid * xbar
                y_product = intensity_grid * ybar
                z_product = intensity_grid * zbar
                integrate = getattr(np, "trapezoid", None) or getattr(np, "trapz")
                X = float(integrate(x_product, wl_grid))
                Y = float(integrate(y_product, wl_grid))
                Z = float(integrate(z_product, wl_grid))
                total = X + Y + Z
                if not np.isfinite(total) or total <= 0:
                    raise ValueError("La suma X+Y+Z no es valida. Revisa columnas, rango o intensidades.")
                x_coord = X / total
                y_coord = Y / total

                st.markdown("#### Paso 1. Preparar el espectro")
                st.write(
                    "Primero se ordenan los datos, se eliminan valores no numericos y se interpola "
                    "la intensidad al intervalo seleccionado. Con esto todos los calculos usan una "
                    "malla regular de longitudes de onda."
                )
                st.latex(r"I(\lambda) \rightarrow I(380), I(385), I(390), \ldots, I(780)")
                st.write(
                    f"En este ejemplo se usan {len(wl_grid)} puntos entre "
                    f"{float(wl_min):.0f} y {float(wl_max):.0f} nm, cada {float(interval):.0f} nm."
                )

                st.markdown("#### Paso 2. Multiplicar por las funciones CIE")
                st.write(
                    "En cada longitud de onda se multiplica la intensidad del espectro por las tres "
                    "funciones colorimetricas del observador estandar CIE 1931."
                )
                st.latex(r"I(\lambda)\overline{x}(\lambda),\quad I(\lambda)\overline{y}(\lambda),\quad I(\lambda)\overline{z}(\lambda)")

                st.markdown("#### Paso 3. Integrar para obtener X, Y y Z")
                st.write(
                    "La integracion suma el area bajo cada una de esas tres curvas. En la aplicacion "
                    "se usa integracion trapezoidal."
                )
                st.latex(r"X = \int I(\lambda)\overline{x}(\lambda)d\lambda")
                st.latex(r"Y = \int I(\lambda)\overline{y}(\lambda)d\lambda")
                st.latex(r"Z = \int I(\lambda)\overline{z}(\lambda)d\lambda")
                st.write("Reemplazando con los valores calculados para este espectro:")
                st.latex(fr"X \approx {X:.4g}")
                st.latex(fr"Y \approx {Y:.4g}")
                st.latex(fr"Z \approx {Z:.4g}")

                st.markdown("#### Paso 4. Normalizar a coordenadas cromaticas")
                st.write(
                    "Como X, Y y Z todavia dependen de la intensidad total, se dividen por "
                    "la suma X+Y+Z para quedarse solo con la cromaticidad."
                )
                st.latex(fr"X + Y + Z = {X:.4g} + {Y:.4g} + {Z:.4g} = {total:.4g}")
                st.latex(fr"x = \frac{{X}}{{X+Y+Z}} = \frac{{{X:.4g}}}{{{total:.4g}}} = {x_coord:.4f}")
                st.latex(fr"y = \frac{{Y}}{{X+Y+Z}} = \frac{{{Y:.4g}}}{{{total:.4g}}} = {y_coord:.4f}")

                r1, r2, r3, r4, r5 = st.columns(5)
                r1.metric("X", f"{X:.4g}")
                r2.metric("Y", f"{Y:.4g}")
                r3.metric("Z", f"{Z:.4g}")
                r4.metric("x", f"{x_coord:.4f}")
                r5.metric("y", f"{y_coord:.4f}")

                plot_col, cie_col = st.columns(2)
                with plot_col:
                    st.markdown("#### Espectro de emision")
                    fig_example, ax_example = plt.subplots(figsize=(5, 4))
                    ax_example.plot(wl_grid, intensity_grid, color="#1f77b4", label="Espectro interpolado")
                    ax_example.set_xlabel("Longitud de onda (nm)")
                    ax_example.set_ylabel("Intensidad (u.a.)")
                    ax_example.grid(alpha=0.25)
                    ax_example.legend(loc="best")
                    show_and_close(fig_example)

                with cie_col:
                    st.markdown("#### Punto en el diagrama CIE")
                    fig_cie_example, ax_cie_example = plt.subplots(figsize=(5, 4))
                    try:
                        try:
                            fig_cie_example, ax_cie_example = colour.plotting.plot_chromaticity_diagram_CIE1931(show=False)
                        except TypeError:
                            fig_cie_example, ax_cie_example = colour.plotting.plot_chromaticity_diagram_CIE1931(standalone=False)
                    except Exception:
                        locus_mask = (CMF_WLS >= 380) & (CMF_WLS <= 780)
                        lx = CMF_X[locus_mask]
                        ly = CMF_Y[locus_mask]
                        lz = CMF_Z[locus_mask]
                        den = lx + ly + lz
                        ax_cie_example.plot(lx / den, ly / den, color="black", linewidth=1.0)
                        ax_cie_example.set_xlim(0, 0.8)
                        ax_cie_example.set_ylim(0, 0.9)
                    ax_cie_example.scatter([x_coord], [y_coord], color="#d62728", s=90, zorder=5, label="Muestra")
                    ax_cie_example.text(x_coord + 0.01, y_coord, f"x={x_coord:.3f}\ny={y_coord:.3f}", fontsize=8)
                    ax_cie_example.set_xlabel("x")
                    ax_cie_example.set_ylabel("y")
                    try:
                        ax_cie_example.legend(loc="best", fontsize=8)
                    except Exception:
                        pass
                    show_and_close(fig_cie_example)

                st.markdown("#### Tabla del calculo")
                calc_df = pd.DataFrame({
                    "wavelength_nm": wl_grid,
                    "I(lambda)": intensity_grid,
                    "xbar(lambda)": xbar,
                    "ybar(lambda)": ybar,
                    "zbar(lambda)": zbar,
                    "I*xbar": x_product,
                    "I*ybar": y_product,
                    "I*zbar": z_product,
                })
                st.dataframe(calc_df.head(30), use_container_width=True)
                st.caption(
                    "La tabla muestra las primeras filas. Las columnas I*xbar, I*ybar e I*zbar "
                    "son las que se integran numericamente para obtener X, Y y Z."
                )
            except Exception as e:
                st.error(f"No se pudo calcular el ejemplo: {e}")
        else:
            st.info("Sube un CSV o XLSX para ver el calculo aplicado a tus propios datos.")

    elif topic == "Visor de espectros":
        st.markdown("### Tipos de espectro")
        st.write(
            "Un espectro de absorcion muestra cuanta luz absorbe la muestra en cada longitud de onda "
            "(util para determinar concentracion con la ley de Beer-Lambert). Un espectro de emision "
            "muestra la intensidad de luz emitida tras excitar la muestra a una longitud de onda fija. "
            "Un espectro de excitacion muestra que longitudes de onda de excitacion producen mas emision "
            "a una longitud de onda de deteccion fija."
        )
        st.markdown("### Solucion vs solido")
        st.write(
            "En solucion, la muestra esta disuelta en un solvente y el entorno quimico es mas homogeneo; "
            "los efectos de filtro interno y las interacciones soluto-solvente son relevantes. En solido "
            "(polvo, pastilla, pelicula), el empaquetamiento cristalino y los efectos de dispersion de luz "
            "pueden ensanchar o desplazar las bandas frente a la misma especie en solucion."
        )
        st.markdown("### Metricas que calcula el visor")
        st.write(
            "Para cada espectro cargado se reportan: longitud de onda del pico maximo, intensidad en el "
            "pico, area integrada (regla trapezoidal) y FWHM (ancho a media altura), calculado como la "
            "distancia entre los primeros puntos que cruzan la mitad de la intensidad maxima a cada lado del pico."
        )
        st.markdown("### Normalizacion")
        st.write(
            "'Max = 1' escala el espectro para que su punto mas alto valga 1, util para comparar formas de "
            "banda entre muestras de intensidad muy distinta. 'Area = 1' escala para que el area bajo la "
            "curva sea 1, util cuando interesa comparar distribucion espectral independientemente de la "
            "intensidad total emitida o absorbida."
        )

    else:  # Rendimiento cuantico
        st.markdown("### Que es el rendimiento cuantico")
        st.write(
            "El rendimiento cuantico de fluorescencia (Phi) es la fraccion de fotones absorbidos que la "
            "muestra reemite como fotones de fluorescencia. El metodo relativo compara el area de emision "
            "integrada de la muestra contra la de una referencia de Phi conocido, medidas bajo la misma "
            "longitud de onda de excitacion y en condiciones comparables."
        )
        st.latex(r"\Phi_x = \Phi_{ref}\left(\frac{I_x}{I_{ref}}\right)\left(\frac{A_{ref}}{A_x}\right)\left(\frac{n_x^2}{n_{ref}^2}\right)")
        st.write(
            "Donde I es el area de emision integrada, A es la absorbancia a la longitud de onda de "
            "excitacion, y n es el indice de refraccion del solvente de cada muestra."
        )
        st.markdown("### Buenas practicas")
        st.write(
            "Mantener la absorbancia baja (tipicamente A < 0.05-0.1 en la longitud de onda de excitacion) "
            "reduce los efectos de filtro interno y de reabsorcion, que de otro modo subestiman el area de "
            "emision real. Tambien conviene medir varias diluciones y graficar area integrada contra "
            "absorbancia: la pendiente de esa recta (en vez de un solo punto) da un Phi mas confiable."
        )

    st.stop()

# ============================================================
# Pagina: Visor de espectros
# ============================================================
if st.session_state["active_page"] == "Visor de espectros":
    st.title("Visor de espectros")
    st.caption("Compara espectros de absorcion, emision o excitacion en solucion y solido.")

    viewer_files = st.file_uploader(
        "Sube archivos CSV o XLSX",
        type=["csv", "xlsx"],
        accept_multiple_files=True,
        key="viewer_files",
    )

    viewer_rows = []
    viewer_spectra = []
    if viewer_files:
        for i, f in enumerate(viewer_files):
            raw = f.getvalue()
            try:
                if f.name.lower().endswith(".xlsx"):
                    sheet_names = list_excel_sheets(raw)
                    selected_sheet = st.selectbox(f"Hoja para {f.name}", options=sheet_names, key=f"viewer_sheet_{i}")
                    df_view = load_excel_sheet(raw, selected_sheet)
                    base_name = f"{f.name} - {selected_sheet}"
                else:
                    df_view = load_csv_df(raw)
                    base_name = f.name

                wl_guess, int_guess = guess_columns(df_view)
                cols = list(df_view.columns)
                with st.expander(f"Configurar {base_name}", expanded=i == 0):
                    c1, c2, c3 = st.columns(3)
                    with c1:
                        label = st.text_input("Etiqueta", value=base_name, key=f"viewer_label_{i}")
                        wl_col = st.selectbox("Columna longitud de onda", options=cols, index=cols.index(wl_guess) if wl_guess in cols else 0, key=f"viewer_wl_{i}")
                    with c2:
                        int_col = st.selectbox("Columna intensidad", options=cols, index=cols.index(int_guess) if int_guess in cols else min(1, len(cols) - 1), key=f"viewer_int_{i}")
                        spectrum_type = st.selectbox("Tipo", options=["Absorcion", "Emision", "Excitacion"], key=f"viewer_type_{i}")
                    with c3:
                        sample_state = st.selectbox("Medio", options=["Solucion", "Solido"], key=f"viewer_state_{i}")
                        line_color = st.color_picker("Color", value=["#1f77b4", "#d62728", "#2ca02c", "#9467bd", "#ff7f0e"][i % 5], key=f"viewer_color_{i}")

                    c4, c5, c6 = st.columns(3)
                    with c4:
                        normalize = st.selectbox("Normalizacion", options=["None", "Max = 1", "Area = 1"], key=f"viewer_norm_{i}")
                    with c5:
                        wl_min_v = st.number_input("Min nm", min_value=100, max_value=10000, value=200, key=f"viewer_min_{i}")
                    with c6:
                        wl_max_v = st.number_input("Max nm", min_value=100, max_value=10000, value=900, key=f"viewer_max_{i}")

                wl, intensity = filter_and_normalize(df_view[wl_col], df_view[int_col], wl_min_v, wl_max_v, normalize)
                peak_wl, peak_intensity, area, fwhm = spectrum_metrics(wl, intensity)
                viewer_spectra.append({
                    "label": label,
                    "wl": wl,
                    "intensity": intensity,
                    "color": line_color,
                })
                viewer_rows.append({
                    "Etiqueta": label,
                    "Tipo": spectrum_type,
                    "Medio": sample_state,
                    "Pico_nm": peak_wl,
                    "Intensidad_pico": peak_intensity,
                    "Area": area,
                    "FWHM_nm": fwhm,
                })
            except Exception as e:
                st.error(f"{f.name}: {e}")

    if viewer_spectra:
        left, right = st.columns([2, 1], gap="large")
        with left:
            fig_v, ax_v = plt.subplots(figsize=(8, 4.8))
            for s in viewer_spectra:
                ax_v.plot(s["wl"], s["intensity"], label=s["label"], color=s["color"], linewidth=1.8)
            ax_v.set_xlabel("Longitud de onda (nm)")
            ax_v.set_ylabel("Intensidad / Absorbancia")
            ax_v.grid(alpha=0.25)
            try:
                ax_v.legend(fontsize=8, loc="best")
            except Exception:
                pass
            show_and_close(fig_v)
        with right:
            st.markdown("### Resumen")
            viewer_df = pd.DataFrame(viewer_rows)
            st.dataframe(viewer_df, use_container_width=True, hide_index=True)
            csv_viewer = io.StringIO()
            viewer_df.to_csv(csv_viewer, index=False)
            st.download_button("Descargar resumen CSV", data=csv_viewer.getvalue(), file_name="spectra_viewer_summary.csv", mime="text/csv")
    else:
        st.info("Sube uno o mas espectros para iniciar la visualizacion.")
    st.stop()

# ============================================================
# Pagina: Rendimiento cuantico
# ============================================================
if st.session_state["active_page"] == "Rendimiento cuantico":
    st.title("Rendimiento cuantico")
    st.caption("Calcula el rendimiento cuantico de fluorescencia por metodo relativo o absoluto.")

    show_explanations = st.toggle(
        "Mostrar explicaciones conceptuales",
        value=True, key="show_explanations",
        help="Actívalo para ver, junto a cada paso, una frase sobre qué significa esa cuenta. "
             "Desactívalo para un desarrollo mas compacto, solo con la matematica.",
    )
    st.caption(
        "Los reportes en PDF y Word siempre incluyen el desarrollo completo con explicaciones, "
        "sin importar esta opcion."
    )

    tab_rel, tab_abs = st.tabs(["Rendimiento cuantico relativo", "Rendimiento cuantico absoluto"])

    # ======================================================
    # Seccion: rendimiento cuantico RELATIVO
    # ======================================================
    with tab_rel:
        st.write(
            "Usa areas integradas de emision, absorbancias a la longitud de onda de excitacion "
            "e indices de refraccion. Mantener absorbancias bajas ayuda a reducir errores por filtro interno."
        )

        col_in, col_out = st.columns([1, 1.2], gap="large")

        with col_in:
            st.markdown("#### Datos de entrada")

            st.markdown("##### 🧪 Muestra")
            st.caption("Datos de tu compuesto problema.")
            sample_area = area_input(
                "Area integrada de emision (muestra)", "rel_sample_area", default_value=1.0,
                help_text="Ix: area bajo la curva del espectro de emision de tu muestra.",
            )
            sample_abs = absorbance_input(
                "Absorbancia/reflectancia a la excitacion (muestra)", "rel_sample_abs", default_value=0.05,
                help_text="Ax: cuanta luz de excitacion absorbe la muestra. Valores bajos (<0.1) evitan errores de filtro interno.",
            )
            sample_n = st.number_input(
                "Indice refraccion muestra", min_value=1.0, value=1.333, format="%.6f", key="rel_sample_n",
                help="nx: indice de refraccion del solvente de la muestra (agua \u2248 1.333, etanol \u2248 1.36).",
            )

            st.markdown("---")
            st.markdown("##### 📖 Referencia")
            st.caption("Compuesto de Φ ya conocido, usado como patron de comparacion.")
            ref_phi = st.number_input(
                "Phi referencia", min_value=0.0, max_value=1.0, value=0.55, format="%.6f", key="rel_ref_phi",
                help="\u03a6ref: rendimiento cuantico ya conocido de la referencia, tomado de la literatura.",
            )
            ref_area = area_input(
                "Area integrada de emision (referencia)", "rel_ref_area", default_value=1.0,
                help_text="Iref: area bajo la curva del espectro de emision de la referencia.",
                embedded_example=QUININE_SULFATE_EXAMPLE,
            )
            ref_abs = absorbance_input(
                "Absorbancia/reflectancia a la excitacion (referencia)", "rel_ref_abs", default_value=0.05,
                help_text="Aref: absorbancia de la referencia a la misma longitud de onda de excitacion.",
                embedded_example=QUININE_SULFATE_ABS_EXAMPLE,
            )
            ref_n = st.number_input(
                "Indice refraccion referencia", min_value=1.0, value=1.333, format="%.6f", key="rel_ref_n",
                help="nref: indice de refraccion del solvente de la referencia.",
            )

        rel_data = {
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "sample_area": sample_area, "sample_abs": sample_abs, "sample_n": sample_n,
            "ref_phi": ref_phi, "ref_area": ref_area, "ref_abs": ref_abs, "ref_n": ref_n,
        }
        phi, rel_steps = _relative_qy_steps(rel_data)

        with col_out:
            st.markdown("#### Formula y desarrollo")
            st.latex(r"\Phi_x = \Phi_{ref}\left(\frac{I_x}{I_{ref}}\right)\left(\frac{A_{ref}}{A_x}\right)\left(\frac{n_x^2}{n_{ref}^2}\right)")

            ratio_area = sample_area / ref_area
            ratio_abs = ref_abs / sample_abs
            ratio_n2 = (sample_n ** 2) / (ref_n ** 2)

            st.markdown("**Paso 1.** Razon de areas de emision integradas:")
            if show_explanations:
                st.caption(rel_steps[1][2])
            st.latex(fr"\frac{{I_x}}{{I_{{ref}}}} = \frac{{{sample_area:.4g}}}{{{ref_area:.4g}}} = {ratio_area:.4g}")

            st.markdown("**Paso 2.** Razon de absorbancias a la longitud de onda de excitacion:")
            if show_explanations:
                st.caption(rel_steps[2][2])
            st.latex(fr"\frac{{A_{{ref}}}}{{A_x}} = \frac{{{ref_abs:.4g}}}{{{sample_abs:.4g}}} = {ratio_abs:.4g}")

            st.markdown("**Paso 3.** Razon de indices de refraccion al cuadrado:")
            if show_explanations:
                st.caption(rel_steps[3][2])
            st.latex(fr"\frac{{n_x^2}}{{n_{{ref}}^2}} = \frac{{{sample_n:.4g}^2}}{{{ref_n:.4g}^2}} = {ratio_n2:.4g}")

            st.markdown("**Paso 4.** Sustituir y multiplicar todo junto con \u03a6ref:")
            if show_explanations:
                st.caption(rel_steps[4][2])
            st.latex(
                fr"\Phi_x = {ref_phi:.4g}\left({ratio_area:.4g}\right)\left({ratio_abs:.4g}\right)\left({ratio_n2:.4g}\right) = {phi:.4g}"
            )
            if show_explanations:
                st.caption(rel_steps[5][2])

            st.metric("Rendimiento cuantico relativo de la muestra", f"{phi:.4f}", f"{phi * 100:.2f}%")

            if phi > 1:
                st.warning(
                    "El resultado es mayor que 1, lo cual no es fisicamente posible (nunca se puede emitir "
                    "mas fotones de los que se absorben). Causas tipicas:\n"
                    "- **Efecto de filtro interno o reabsorcion**: absorbancias demasiado altas (>0.1) distorsionan el area de emision.\n"
                    "- **Rango de integracion inconsistente**: revisa que el area de muestra y referencia cubran el mismo rango espectral.\n"
                    "- **\u03a6ref incorrecto**: verifica el valor de referencia contra la literatura, incluyendo el solvente y la longitud de onda usados.\n"
                    "- **Linea base sin corregir**: una linea base desplazada infla el area integrada."
                )
            elif show_explanations:
                if phi < 0.1:
                    nivel = "bajo"
                    comentario = "tipico de emisores con mucha desactivacion no radiativa (perdida de energia como calor)."
                elif phi < 0.5:
                    nivel = "moderado"
                    comentario = "un rango comun para muchos complejos organicos y de lantanidos con ligandos organicos."
                else:
                    nivel = "alto"
                    comentario = "tipico de buenos fluoroforos organicos (p. ej. rodaminas, fluoresceina)."
                st.caption(
                    f"Orientativo: un \u03a6 {nivel} como este es {comentario} Esto es solo una guia general, "
                    "no una regla estricta; el valor esperado depende mucho del sistema quimico especifico."
                )

            st.markdown("#### Exportar resultado")
            col_pdf, col_docx = st.columns(2)
            with col_pdf:
                if _HAS_REPORTLAB:
                    try:
                        pdf_bytes = build_relative_qy_pdf(rel_data)
                        st.download_button("Descargar PDF", data=pdf_bytes, file_name="rendimiento_cuantico_relativo.pdf",
                                           mime="application/pdf", use_container_width=True, key="rel_pdf")
                    except Exception as e:
                        st.error(f"No se pudo generar el PDF: {e}")
                else:
                    st.info("Para exportar a PDF instala la libreria 'reportlab' (pip install reportlab).")
            with col_docx:
                if _HAS_DOCX:
                    try:
                        docx_bytes = build_relative_qy_docx(rel_data)
                        st.download_button("Descargar Word (.docx)", data=docx_bytes, file_name="rendimiento_cuantico_relativo.docx",
                                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                           use_container_width=True, key="rel_docx")
                    except Exception as e:
                        st.error(f"No se pudo generar el documento Word: {e}")
                else:
                    st.info("Para exportar a Word instala la libreria 'python-docx' (pip install python-docx).")

    # ======================================================
    # Seccion: rendimiento cuantico ABSOLUTO
    # ======================================================
    with tab_abs:
        st.write(
            "Metodo de esfera integradora (de Mello, Wittmann y Friend, 1997): no necesita una "
            "referencia externa, pero requiere cuatro areas integradas medidas con el mismo montaje."
        )
        st.caption(
            "La = area de excitacion dispersada (referencia/blanco, sin excitar directamente la muestra). "
            "Lc = area de excitacion dispersada con la muestra en el haz directo. "
            "Ea = area de emision con excitacion indirecta. Ec = area de emision con excitacion directa."
        )

        col_in2, col_out2 = st.columns([1, 1.2], gap="large")

        with col_in2:
            st.markdown("#### Datos de entrada")
            if st.button("Cargar ejemplo", key="abs_load_example", use_container_width=True):
                st.session_state["abs_La_mode"] = "Manual"
                st.session_state["abs_La_manual"] = 1.0
                st.session_state["abs_Lc_mode"] = "Manual"
                st.session_state["abs_Lc_manual"] = 0.35
                st.session_state["abs_Ea_mode"] = "Manual"
                st.session_state["abs_Ea_manual"] = 0.05
                st.session_state["abs_Ec_mode"] = "Manual"
                st.session_state["abs_Ec_manual"] = 0.42
                st.rerun()

            st.markdown("##### 🔦 Excitacion")
            st.caption("Area del pico de excitacion dispersada.")
            L_a = area_input("La \u2014 referencia/blanco", "abs_La", default_value=1.0,
                             help_text="Area del pico de excitacion dispersada, midiendo solo el blanco o la referencia.")
            L_c = area_input("Lc \u2014 con la muestra", "abs_Lc", default_value=0.5,
                             help_text="Area del mismo pico de excitacion, ahora con la muestra puesta en el haz directo.")

            st.markdown("---")
            st.markdown("##### ✨ Emision")
            st.caption("Area del pico de emision de la muestra.")
            E_a = area_input("Ea \u2014 excitacion indirecta", "abs_Ea", default_value=0.1,
                             help_text="Area de emision cuando la muestra se excita solo indirectamente (luz dispersada dentro de la esfera).")
            E_c = area_input("Ec \u2014 excitacion directa", "abs_Ec", default_value=0.6,
                             help_text="Area de emision cuando la muestra se excita directamente con el haz.")

        abs_data = {
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "L_a": L_a, "L_c": L_c, "E_a": E_a, "E_c": E_c,
        }

        with col_out2:
            st.markdown("#### Formula y desarrollo")
            st.latex(r"A = 1 - \frac{L_c}{L_a}")
            st.latex(r"\Phi_{abs} = \frac{E_c - (1-A)\,E_a}{A\,L_a}")

            if L_c > L_a:
                st.error("Lc no puede ser mayor que La (la muestra no puede dispersar/transmitir mas luz de la que llega).")
            else:
                phi_abs, A, abs_steps = _absolute_qy_steps(abs_data)
                lc_la_ratio = L_c / L_a

                st.markdown("**Paso 1.** Fraccion de luz absorbida por la muestra:")
                if show_explanations:
                    st.caption(abs_steps[0][2])
                st.latex(fr"A = 1 - \frac{{L_c}}{{L_a}} = 1 - \frac{{{L_c:.4g}}}{{{L_a:.4g}}} = 1 - {lc_la_ratio:.4g} = {A:.4g}")

                st.markdown("**Paso 2.** Termino de emision indirecta corregido:")
                if show_explanations:
                    st.caption(abs_steps[2][2])
                st.latex(fr"(1-A)\,E_a = (1 - {A:.4g})\times {E_a:.4g} = {(1 - A) * E_a:.4g}")

                st.markdown("**Paso 3.** Numerador (emision directa menos la indirecta corregida):")
                if show_explanations:
                    st.caption(abs_steps[3][2])
                st.latex(fr"E_c - (1-A)E_a = {E_c:.4g} - {(1 - A) * E_a:.4g} = {E_c - (1 - A) * E_a:.4g}")

                st.markdown("**Paso 4.** Denominador:")
                if show_explanations:
                    st.caption(abs_steps[4][2])
                st.latex(fr"A \times L_a = {A:.4g} \times {L_a:.4g} = {A * L_a:.4g}")

                st.markdown("**Paso 5.** Dividir para obtener el rendimiento cuantico absoluto:")
                if show_explanations:
                    st.caption(abs_steps[5][2])
                st.latex(
                    fr"\Phi_{{abs}} = \frac{{{E_c - (1 - A) * E_a:.4g}}}{{{A * L_a:.4g}}} = {phi_abs:.4g}"
                )

                st.metric("Rendimiento cuantico absoluto de la muestra", f"{phi_abs:.4f}", f"{phi_abs * 100:.2f}%")

                if phi_abs > 1:
                    st.warning(
                        "El resultado es mayor que 1, lo cual no es fisicamente posible. Causas tipicas:\n"
                        "- **Mala separacion entre el pico de excitacion y el de emision** al integrar las areas.\n"
                        "- **Geometria inconsistente** entre las mediciones de La/Lc y Ea/Ec (deben usar el mismo montaje).\n"
                        "- **Muestra que dispersa mucha luz** (solidos/polvos), lo que distorsiona el pico de excitacion."
                    )
                elif show_explanations:
                    if phi_abs < 0.1:
                        nivel = "bajo"
                        comentario = "tipico de emisores con mucha desactivacion no radiativa."
                    elif phi_abs < 0.5:
                        nivel = "moderado"
                        comentario = "un rango comun para complejos de lantanidos y materiales solidos luminiscentes."
                    else:
                        nivel = "alto"
                        comentario = "tipico de buenos emisores solidos o fosforos comerciales."
                    st.caption(
                        f"Orientativo: un \u03a6 {nivel} como este es {comentario} Esto es solo una guia general, "
                        "no una regla estricta."
                    )

                st.markdown("#### Exportar resultado")
                col_pdf2, col_docx2 = st.columns(2)
                with col_pdf2:
                    if _HAS_REPORTLAB:
                        try:
                            pdf_bytes_abs = build_absolute_qy_pdf(abs_data)
                            st.download_button("Descargar PDF", data=pdf_bytes_abs, file_name="rendimiento_cuantico_absoluto.pdf",
                                               mime="application/pdf", use_container_width=True, key="abs_pdf")
                        except Exception as e:
                            st.error(f"No se pudo generar el PDF: {e}")
                    else:
                        st.info("Para exportar a PDF instala la libreria 'reportlab' (pip install reportlab).")
                with col_docx2:
                    if _HAS_DOCX:
                        try:
                            docx_bytes_abs = build_absolute_qy_docx(abs_data)
                            st.download_button("Descargar Word (.docx)", data=docx_bytes_abs, file_name="rendimiento_cuantico_absoluto.docx",
                                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                               use_container_width=True, key="abs_docx")
                        except Exception as e:
                            st.error(f"No se pudo generar el documento Word: {e}")
                    else:
                        st.info("Para exportar a Word instala la libreria 'python-docx' (pip install python-docx).")

    st.stop()

# ============================================================
# Pagina: Analisis CIE 1931 (principal)
# ============================================================
st.title("CIE 1931 - Coordenadas de cromaticidad a partir de espectros de emision")

# ----------------- Sidebar (parametros de esta pagina) -----------------
st.sidebar.header("Parametros globales")
wl_min_default = st.sidebar.number_input("Longitud de onda minima (nm)", min_value=200, max_value=10000, value=380)
wl_max_default = st.sidebar.number_input("Longitud de onda maxima (nm)", min_value=200, max_value=10000, value=780)
interp_interval_default = st.sidebar.selectbox("Intervalo de interpolacion (nm)", options=[1, 2, 5, 10], index=2)
dpi_save = st.sidebar.number_input("DPI para exportar imagenes", min_value=72, max_value=1200, value=600)

# ---- Personalizacion del diagrama ----
st.sidebar.markdown("---")
st.sidebar.subheader("Personalizacion del diagrama")

plot_title = st.sidebar.text_input("Titulo del grafico", value="Diagrama de cromaticidad CIE 1931")
x_axis_label = st.sidebar.text_input("Etiqueta eje X", value="Coordenada de cromaticidad x")
y_axis_label = st.sidebar.text_input("Etiqueta eje Y", value="Coordenada de cromaticidad y")

title_color = st.sidebar.color_picker("Color del titulo", "#000000")
axes_color = st.sidebar.color_picker("Color de ejes y marcas", "#000000")
locus_label_color = st.sidebar.color_picker("Color de numeros de longitud de onda", "#000000")

title_font_family = st.sidebar.selectbox("Fuente del titulo", options=["sans-serif", "serif", "monospace"], index=0)
title_font_size = st.sidebar.number_input("Tamano de fuente del titulo", min_value=8, max_value=48, value=14)

tick_font_family = st.sidebar.selectbox("Fuente de marcas", options=["sans-serif", "serif", "monospace"], index=0)
tick_font_size = st.sidebar.number_input("Tamano de marcas", min_value=6, max_value=24, value=10)

locus_numbers_font_family = st.sidebar.selectbox("Fuente de numeros (locus)", options=["sans-serif", "serif", "monospace"], index=0)
locus_numbers_font_size = st.sidebar.number_input("Tamano de numeros (locus)", min_value=6, max_value=24, value=8)

axes_linewidth = st.sidebar.slider("Grosor de ejes", min_value=0.5, max_value=5.0, value=1.0, step=0.1)

show_point_labels = st.sidebar.checkbox("Mostrar etiquetas junto a cada punto", value=True)

st.sidebar.markdown("---")
st.sidebar.subheader("Leyenda")
legend_loc = st.sidebar.selectbox("Ubicacion", options=[
    "best", "upper right", "upper left", "lower left", "lower right",
    "right", "center left", "center right", "lower center", "upper center", "center"
], index=0)
legend_font_size = st.sidebar.number_input("Tamano de fuente", min_value=6, max_value=24, value=8)
legend_ncols = st.sidebar.slider("Columnas", min_value=1, max_value=4, value=1)
legend_box_color = st.sidebar.color_picker("Fondo", "#FFFFFF")
legend_box_alpha = st.sidebar.slider("Opacidad de fondo", min_value=0.0, max_value=1.0, value=0.85)
legend_box_linewidth = st.sidebar.slider("Grosor del borde", min_value=0.0, max_value=4.0, value=0.8, step=0.1)

# ---- Longitud dominante / pureza ----
st.sidebar.markdown("---")
st.sidebar.subheader("Longitud de onda dominante / pureza")

wp_mode = st.sidebar.selectbox("Punto blanco", ["E (0.3333, 0.3333)", "D65 (0.3127, 0.3290)", "Personalizado"], index=0)
if wp_mode.startswith("E"):
    WHITE_POINT = (0.3333, 0.3333)
elif wp_mode.startswith("D65"):
    WHITE_POINT = (0.3127, 0.3290)
else:
    wp_x = st.sidebar.number_input("Punto blanco x", min_value=0.0, max_value=1.0, value=0.3333, step=0.0001, format="%.4f")
    wp_y = st.sidebar.number_input("Punto blanco y", min_value=0.0, max_value=1.0, value=0.3333, step=0.0001, format="%.4f")
    WHITE_POINT = (float(wp_x), float(wp_y))

# ----------------- Carga de archivos -----------------
st.markdown("### Subir datos")
st.caption("CSV: uno o mas archivos. XLSX: uno o mas archivos (se procesan todas las hojas).")

csv_files = st.file_uploader("Sube uno o mas archivos CSV", type=["csv"], accept_multiple_files=True)
xlsx_files = st.file_uploader("Sube uno o mas archivos Excel (.xlsx)", type=["xlsx"], accept_multiple_files=True)

# ----------------- Construccion de datasets -----------------
datasets = []

if csv_files:
    for i, f in enumerate(csv_files):
        raw = uploaded_to_bytes(f)
        if not raw:
            continue
        try:
            df = load_csv_df(raw)
            datasets.append({"id": f"csv_{i}", "name": f.name, "df": df})
        except Exception as e:
            st.error(f"Error al leer CSV {f.name}: {e}")

if xlsx_files:
    for j, f in enumerate(xlsx_files):
        raw = uploaded_to_bytes(f)
        if not raw:
            continue
        try:
            sheets = list_excel_sheets(raw)
        except Exception as e:
            st.error(f"Error al leer Excel {f.name}: {e}")
            continue
        for k, sh in enumerate(sheets):
            try:
                df = load_excel_sheet(raw, sh)
                datasets.append({"id": f"xlsx_{j}_{k}", "name": f"{f.name} - {sh}", "df": df})
            except Exception as e:
                st.error(f"Error al leer {f.name} | hoja {sh}: {e}")

# ----------------- Configuracion por dataset -----------------
def init_cfg(ds):
    df = ds["df"]
    cols = list(df.columns)
    wl_guess, int_guess = guess_columns(df)
    return {
        "label": ds["name"],
        "color": "#000000",
        "spectrum_color": "#1f77b4",
        "marker": "o",
        "size": 120,
        "wl_col": wl_guess if wl_guess in cols else (cols[0] if cols else None),
        "int_col": int_guess if int_guess in cols else (cols[1] if len(cols) > 1 else (cols[0] if cols else None)),
        "wl_min": int(wl_min_default),
        "wl_max": int(wl_max_default),
        "interp": int(interp_interval_default),
        "baseline_subtract_min": False,
        "clip_negative": False,
        "smooth_method": "None",
        "smooth_window": 11,
        "smooth_poly": 3,
        "normalize": "None",
    }


def get_cfg(ds):
    key = f"cfg_{ds['id']}"
    if key not in st.session_state:
        st.session_state[key] = init_cfg(ds)
    cfg = st.session_state[key]
    cols = list(ds["df"].columns)
    if cfg.get("wl_col") not in cols and cols:
        cfg["wl_col"] = cols[0]
    if cfg.get("int_col") not in cols and cols:
        cfg["int_col"] = cols[1] if len(cols) > 1 else cols[0]
    if "spectrum_color" not in cfg:
        cfg["spectrum_color"] = cfg.get("color", "#1f77b4")
    st.session_state[key] = cfg
    return cfg


def render_cfg_form(ds, cfg, key_prefix="dlg_"):
    """Formulario de configuracion por muestra.
    key_prefix evita colisiones de keys de widgets entre dialogos."""
    ds_id = ds["id"]
    keyp = f"{key_prefix}{ds_id}_"
    cols = list(ds["df"].columns)

    st.markdown("Columnas")
    wl_col = st.selectbox("Columna de longitud de onda", options=cols,
                          index=cols.index(cfg["wl_col"]) if cfg["wl_col"] in cols else 0,
                          key=f"{keyp}wlcol")
    int_default_index = cols.index(cfg["int_col"]) if cfg["int_col"] in cols else min(1, len(cols) - 1)
    int_col = st.selectbox("Columna de intensidad", options=cols,
                           index=int_default_index,
                           key=f"{keyp}intcol")

    st.markdown("Estilo del punto")
    label = st.text_input("Etiqueta", value=cfg["label"], key=f"{keyp}label")
    color = st.color_picker("Color", value=cfg["color"], key=f"{keyp}color")
    spectrum_color = st.color_picker("Color de la linea del espectro", value=cfg["spectrum_color"], key=f"{keyp}spectrum_color")
    marker_opts = ['o', 's', '^', 'x', 'D', '*', 'v']
    marker = st.selectbox("Marcador", options=marker_opts,
                          index=marker_opts.index(cfg["marker"]) if cfg["marker"] in marker_opts else 0,
                          key=f"{keyp}marker")
    size = st.slider("Tamano (s)", min_value=20, max_value=500, value=int(cfg["size"]), key=f"{keyp}size")

    st.markdown("Rango de calculo")
    wl_min = st.number_input("Longitud de onda minima (nm)", min_value=200, max_value=10000, value=int(cfg["wl_min"]), key=f"{keyp}wlmin")
    wl_max = st.number_input("Longitud de onda maxima (nm)", min_value=200, max_value=10000, value=int(cfg["wl_max"]), key=f"{keyp}wlmax")
    interp_opts = [1, 2, 5, 10]
    interp = st.selectbox("Intervalo (nm)", options=interp_opts,
                          index=interp_opts.index(int(cfg["interp"])) if int(cfg["interp"]) in interp_opts else 1,
                          key=f"{keyp}interp")

    st.markdown("Preprocesamiento")
    baseline_subtract_min = st.checkbox("Linea base: restar minimo (desplazar a 0)", value=bool(cfg["baseline_subtract_min"]), key=f"{keyp}base")
    clip_negative = st.checkbox("Recortar intensidades negativas a 0", value=bool(cfg["clip_negative"]), key=f"{keyp}clip")

    smooth_options = ["None", "Moving average"] + (["Savitzky-Golay"] if _HAS_SCIPY else [])
    smooth_method = st.selectbox("Suavizado", options=smooth_options,
                                 index=smooth_options.index(cfg["smooth_method"]) if cfg["smooth_method"] in smooth_options else 0,
                                 key=f"{keyp}smooth")
    smooth_window = st.slider("Ventana (suavizado)", min_value=3, max_value=101, value=int(cfg["smooth_window"]), step=2, key=f"{keyp}win")
    smooth_poly = st.slider("Orden del polinomio (Savitzky-Golay)", min_value=2, max_value=7, value=int(cfg["smooth_poly"]), step=1, key=f"{keyp}poly")
    normalize = st.selectbox("Normalizacion", options=["None", "Max = 1", "Area = 1"],
                             index=["None", "Max = 1", "Area = 1"].index(cfg["normalize"]) if cfg["normalize"] in ["None", "Max = 1", "Area = 1"] else 0,
                             key=f"{keyp}norm")

    return {
        "label": label,
        "color": color,
        "spectrum_color": spectrum_color,
        "marker": marker,
        "size": int(size),
        "wl_col": wl_col,
        "int_col": int_col,
        "wl_min": int(wl_min),
        "wl_max": int(wl_max),
        "interp": int(interp),
        "baseline_subtract_min": bool(baseline_subtract_min),
        "clip_negative": bool(clip_negative),
        "smooth_method": smooth_method,
        "smooth_window": int(smooth_window),
        "smooth_poly": int(smooth_poly),
        "normalize": normalize,
    }


# ----------------- Configuracion por muestra (modal) -----------------
if "open_cfg_id" not in st.session_state:
    st.session_state["open_cfg_id"] = None


def request_open_config(ds_id: str):
    st.session_state["open_cfg_id"] = ds_id
    st.rerun()


def handle_config_dialog(datasets_by_id):
    """Si hay una muestra seleccionada, abre el dialogo y detiene el resto
    del script. Esto evita calcular con configuracion desactualizada en la
    misma ejecucion."""
    ds_id = st.session_state.get("open_cfg_id", None)
    if not ds_id:
        return

    ds = datasets_by_id.get(ds_id)
    if ds is None:
        st.session_state["open_cfg_id"] = None
        return

    cfg = get_cfg(ds)
    title = f"Configurar: {ds['name']}"

    if HAS_DIALOG:
        @DIALOG_DECORATOR(title)
        def _dlg():
            with st.form(key=f"form_{ds_id}"):
                new_cfg = render_cfg_form(ds, cfg, key_prefix="dlg_")
                cA, cB = st.columns(2)
                with cA:
                    submitted = st.form_submit_button("Guardar")
                with cB:
                    cancelled = st.form_submit_button("Cancelar")
            if submitted:
                st.session_state[f"cfg_{ds_id}"] = new_cfg
                st.session_state["open_cfg_id"] = None
                st.rerun()
            if cancelled:
                st.session_state["open_cfg_id"] = None
                st.rerun()

        _dlg()
        st.stop()
    else:
        with st.expander(title, expanded=True):
            new_cfg = render_cfg_form(ds, cfg, key_prefix="dlg_")
            cA, cB = st.columns(2)
            if cA.button("Guardar", key=f"save_{ds_id}"):
                st.session_state[f"cfg_{ds_id}"] = new_cfg
                st.session_state["open_cfg_id"] = None
                st.rerun()
            if cB.button("Cancelar", key=f"cancel_{ds_id}"):
                st.session_state["open_cfg_id"] = None
                st.rerun()
        st.stop()


datasets_by_id = {ds['id']: ds for ds in datasets}
handle_config_dialog(datasets_by_id)

# ----------------- Figura CIE -----------------
fig, ax = plt.subplots(figsize=(7, 7))
try:
    try:
        fig_cie, ax_cie = colour.plotting.plot_chromaticity_diagram_CIE1931(show=False)
    except TypeError:
        fig_cie, ax_cie = colour.plotting.plot_chromaticity_diagram_CIE1931(standalone=False)
    ax.clear()
    plt.close(fig)
    fig, ax = fig_cie, ax_cie
except Exception:
    ax.set_xlim(0, 0.8)
    ax.set_ylim(0, 0.9)

fig.subplots_adjust(left=0.12, right=0.98, top=0.95, bottom=0.12)

for txt in list(ax.texts):
    try:
        txt.set_clip_on(False)
        txt.set_bbox(dict(facecolor="white", edgecolor="none", alpha=0.7, pad=1.0))
    except Exception:
        pass
for lbl in ax.get_xticklabels() + ax.get_yticklabels():
    lbl.set_clip_on(False)

fontprop_title = FontProperties(family=title_font_family, size=title_font_size)
fontprop_ticks = FontProperties(family=tick_font_family, size=tick_font_size)
fontprop_locus = FontProperties(family=locus_numbers_font_family, size=locus_numbers_font_size)

try:
    t = ax.set_title(_safe_mpl_text(plot_title), color=title_color)
    t.set_fontproperties(fontprop_title)
except Exception:
    ax.set_title(_safe_mpl_text(plot_title), color=title_color)

try:
    tx = ax.set_xlabel(_safe_mpl_text(x_axis_label), color=axes_color)
    ty = ax.set_ylabel(_safe_mpl_text(y_axis_label), color=axes_color)
    tx.set_fontproperties(fontprop_ticks)
    ty.set_fontproperties(fontprop_ticks)
except Exception:
    ax.set_xlabel(_safe_mpl_text(x_axis_label), color=axes_color)
    ax.set_ylabel(_safe_mpl_text(y_axis_label), color=axes_color)

for lbl in ax.get_xticklabels() + ax.get_yticklabels():
    try:
        lbl.set_fontproperties(fontprop_ticks)
        lbl.set_color(axes_color)
        lbl.set_clip_on(False)
    except Exception:
        pass

for txt in list(ax.texts):
    try:
        txt.set_clip_on(False)
        txt.set_fontproperties(fontprop_locus)
        txt.set_color(locus_label_color)
    except Exception:
        pass

try:
    for spine in ax.spines.values():
        spine.set_linewidth(axes_linewidth)
except Exception:
    pass

# ----------------- Procesar muestras -----------------
results = []
spectra_to_plot = []

if datasets:
    st.markdown("### Muestras")
    st.caption("Opciones por muestra: boton Configurar (ventana emergente). Los resultados se muestran en el resumen debajo.")

    for ds in datasets:
        cfg = get_cfg(ds)

        c1, c2, c3 = st.columns([6, 2, 4])
        with c1:
            st.write(ds["name"])
        with c2:
            if st.button("Configurar", key=f"btn_cfg_{ds['id']}"):
                request_open_config(ds["id"])
        with c3:
            st.write(f"Longitud de onda: {cfg['wl_min']}-{cfg['wl_max']} nm | interp: {cfg['interp']} nm")

        try:
            df = ds["df"]
            if cfg["wl_min"] >= cfg["wl_max"]:
                raise ValueError("La longitud de onda minima debe ser menor que la maxima.")
            if cfg["wl_min"] < CMF_MIN or cfg["wl_max"] > CMF_MAX:
                raise ValueError(f"El rango debe estar dentro del dominio de las CMF: {CMF_MIN:.0f}-{CMF_MAX:.0f} nm.")
            if cfg["wl_col"] not in df.columns or cfg["int_col"] not in df.columns:
                raise ValueError("Las columnas seleccionadas no existen en este archivo/hoja.")

            wl_grid, intensity_grid = preprocess_spectrum(
                df, cfg["wl_col"], cfg["int_col"],
                cfg["wl_min"], cfg["wl_max"], cfg["interp"],
                clip_negative=cfg["clip_negative"],
                baseline_subtract_min=cfg["baseline_subtract_min"],
                smooth_method=cfg["smooth_method"],
                smooth_window=cfg["smooth_window"],
                smooth_poly=cfg["smooth_poly"],
                normalize_mode=cfg["normalize"],
            )

            x_val, y_val = spectrum_to_xy(wl_grid, intensity_grid)
            wl_dom, purity, wl_kind = dominant_wavelength_and_purity(x_val, y_val, white_point=WHITE_POINT)

            ax.scatter([x_val], [y_val], c=[cfg["color"]], marker=cfg["marker"], s=cfg["size"], label=cfg["label"])
            if show_point_labels:
                ax.text(x_val + 0.008, y_val, cfg["label"],
                        fontsize=tick_font_size, fontfamily=tick_font_family,
                        bbox=dict(facecolor="white", edgecolor="none", alpha=0.7, pad=1.0))

            results.append({
                "Etiqueta": cfg["label"],
                "x": float(x_val),
                "y": float(y_val),
                "Longitud_onda_nm": (float(wl_dom) if np.isfinite(wl_dom) else np.nan),
                "Tipo_longitud_onda": wl_kind,
                "Pureza_excitacion_%": (float(purity) if np.isfinite(purity) else np.nan),
                "wl_min_nm": int(cfg["wl_min"]),
                "wl_max_nm": int(cfg["wl_max"]),
                "interp_nm": int(cfg["interp"]),
                "linea_base_restar_min": bool(cfg["baseline_subtract_min"]),
                "recorte_negativos": bool(cfg["clip_negative"]),
                "metodo_suavizado": cfg["smooth_method"],
                "ventana_suavizado": int(cfg["smooth_window"]),
                "orden_polinomio": int(cfg["smooth_poly"]),
                "normalizacion": cfg["normalize"],
                "columna_wl": cfg["wl_col"],
                "columna_intensidad": cfg["int_col"],
            })

            spectra_to_plot.append({
                "label": cfg["label"],
                "wl": wl_grid,
                "it": intensity_grid,
                "color": cfg.get("spectrum_color", cfg["color"])
            })

            st.success(f"{cfg['label']}: x={x_val:.4f}, y={y_val:.4f} | longitud de onda ({wl_kind})={wl_dom:.1f} nm | pureza={purity:.1f}%")
        except Exception as e:
            st.error(f"{cfg['label']}: {e}")
else:
    st.info("Sube archivos para comenzar.")

# ----------------- Resultados -----------------
if results:
    try:
        legend = ax.legend(loc=legend_loc, fontsize=legend_font_size, ncol=legend_ncols)
        frame = legend.get_frame()
        frame.set_facecolor(legend_box_color)
        frame.set_alpha(legend_box_alpha)
        frame.set_linewidth(legend_box_linewidth)
    except Exception:
        pass

    st.markdown("### Resumen del analisis")
    results_df = pd.DataFrame(results)
    minimal_cols = [
        "Etiqueta",
        "x",
        "y",
        "Longitud_onda_nm",
        "Tipo_longitud_onda",
        "Pureza_excitacion_%",
    ]
    minimal_cols = [c for c in minimal_cols if c in results_df.columns]
    table_df = results_df[minimal_cols].copy()
    for c in ["x", "y"]:
        if c in table_df.columns:
            table_df[c] = table_df[c].map(lambda v: f"{v:.4f}" if np.isfinite(v) else "")
    for c in ["Longitud_onda_nm", "Pureza_excitacion_%"]:
        if c in table_df.columns:
            table_df[c] = table_df[c].map(lambda v: f"{v:.1f}" if np.isfinite(v) else "")

    left_col, right_col = st.columns([1, 1], gap="large")

    with left_col:
        st.markdown("### Diagrama CIE 1931")
        try:
            plt.tight_layout()
        except Exception:
            pass
        show_and_close(fig)

    with right_col:
        st.markdown("### Espectros de emision")
        if spectra_to_plot:
            fig_s, ax_s = plt.subplots(figsize=(5.5, 3.2))
            for s in spectra_to_plot:
                ax_s.plot(s["wl"], s["it"], label=s["label"], color=s["color"], linewidth=1.8)
            ax_s.set_xlabel("Longitud de onda (nm)")
            ax_s.set_ylabel("Intensidad (u.a.)")
            ax_s.tick_params(labelsize=9)
            ax_s.grid(alpha=0.25)
            try:
                ax_s.legend(fontsize=7, loc="best")
            except Exception:
                pass
            show_and_close(fig_s)
        else:
            st.info("No hay espectros para mostrar.")

        st.markdown("### Tabla de coordenadas")
        st.dataframe(table_df, use_container_width=True, hide_index=True)

        with st.expander("Descargas", expanded=False):
            csv_buf = io.StringIO()
            results_df.to_csv(csv_buf, index=False)
            st.download_button("Descargar tabla CSV completa", data=csv_buf.getvalue(), file_name="CIE1931_coordenadas.csv", mime="text/csv")

            st.markdown("**Diagrama CIE 1931**")
            dcol1, dcol2, dcol3 = st.columns(3)
            with dcol1:
                img_buf_png = io.BytesIO()
                fig.savefig(img_buf_png, format="png", dpi=dpi_save, facecolor="white")
                img_buf_png.seek(0)
                st.download_button("Descargar PNG", data=img_buf_png.getvalue(), file_name="CIE1931_diagrama.png", mime="image/png", use_container_width=True)
            with dcol2:
                img_buf_jpg = io.BytesIO()
                fig.savefig(img_buf_jpg, format="jpeg", dpi=dpi_save, facecolor="white")
                img_buf_jpg.seek(0)
                st.download_button("Descargar JPEG", data=img_buf_jpg.getvalue(), file_name="CIE1931_diagrama.jpg", mime="image/jpeg", use_container_width=True)
            with dcol3:
                img_buf_tiff = io.BytesIO()
                fig.savefig(img_buf_tiff, format="tiff", dpi=dpi_save, facecolor="white")
                img_buf_tiff.seek(0)
                st.download_button("Descargar TIFF", data=img_buf_tiff.getvalue(), file_name="CIE1931_diagrama.tiff", mime="image/tiff", use_container_width=True)
else:
    if datasets:
        plt.close(fig)
